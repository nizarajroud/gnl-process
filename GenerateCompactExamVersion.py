#!/usr/bin/env python3
"""Generate compact exam version from PDF with questions."""
import fire
import fitz  # PyMuPDF
import re
import os
import markdown
from pathlib import Path
from dotenv import load_dotenv
from weasyprint import HTML

load_dotenv()


def extract_compact_exam(filename: str, origin: str = 'udemy'):
    """Extract questions, options, and correct answers from PDF.
    
    Args:
        filename: Name of the file without extension (e.g., 'exam' not 'exam.pdf')
        origin: Source of the exam - 'udemy' or 'dojo' (default: 'udemy')
    """
    if origin.lower() not in ['udemy', 'dojo']:
        raise ValueError("origin must be 'udemy' or 'dojo'")
    
    origin = origin.lower()
    
    # Get GNL_PROCESSING_PATH from environment
    gnl_processing_path = os.getenv('GNL_PROCESSING_PATH')
    if not gnl_processing_path:
        raise ValueError("GNL_PROCESSING_PATH not found in .env file")
    
    # Construct input path: GNL_PROCESSING_PATH/../pdf-formatting/pdf/filename.pdf
    base_path = Path(gnl_processing_path).parent
    pdf_path = base_path / "pdf-formatting" / "pdf" / f"{filename}.pdf"
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")
    
    # Construct output path: GNL_PROCESSING_PATH/../Anki-generation/markdown/
    output_dir = base_path / "Anki-generation" / "markdown"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    if origin == 'udemy':
        output_file = extract_udemy_exam(pdf_path, output_dir, filename, base_path)
    else:
        output_file = extract_dojo_exam(pdf_path, output_dir, filename, base_path)
    
    return str(output_file)


def extract_udemy_exam(pdf_path: Path, output_dir: Path, filename: str, base_path: Path):
    """Extract Udemy exam format - read from Word document to preserve formatting."""
    from docx import Document
    
    # Read from Word document instead of PDF to preserve formatting
    word_path = base_path / "pdf-formatting" / "word" / f"{filename}.docx"
    
    if not word_path.exists():
        raise FileNotFoundError(f"Word document not found: {word_path}")
    
    doc = Document(str(word_path))
    
    output_lines = []
    skip_section = False
    prev_was_list_item = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip Udemy-specific patterns
        if any(pattern in text for pattern in ['Ignoré', 'Bonne réponse', 'Sélection correcte', 'Explication générale']):
            continue
        
        # Check if we hit "Correct option(s):" or "Incorrect option(s):" - start skipping entire section
        if re.match(r'^(Correct|Incorrect)\s+options?:', text, re.IGNORECASE):
            skip_section = True
            continue
        
        # Check if we hit a new question - stop skipping
        if re.match(r'^Question\s+\d+:', text):
            skip_section = False
            # Add blank line before question if previous was a list item
            if prev_was_list_item:
                output_lines.append("\n")
                prev_was_list_item = False
        
        # Skip if we're in a section to remove
        if skip_section:
            continue
        
        # Skip empty paragraphs
        if not text:
            continue
        
        # Check if paragraph is bold
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        
        # Check if this is a list item
        if text.startswith('- '):
            if is_bold:
                # Bold list item - extract text after "- " and make it bold
                option_text = text[2:]
                output_lines.append(f"- **{option_text}**\n")
            else:
                output_lines.append(f"{text}\n")
            prev_was_list_item = True
        else:
            # Regular paragraph
            if is_bold:
                output_lines.append(f"**{text}**\n\n")
            else:
                output_lines.append(f"{text}\n\n")
            prev_was_list_item = False
    
    # Save to markdown file
    output_file = output_dir / f"{filename}.md"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(''.join(output_lines))
    
    print(f"Compact exam version saved to: {output_file}")
    
    # Generate PDF from the markdown (save to compact-exam-versions)
    pdf_file = generate_pdf_from_markdown_internal(output_file, filename, base_path)
    print(f"PDF version saved to: {pdf_file}")
    
    return output_file


def extract_dojo_exam(pdf_path: Path, output_dir: Path, filename: str, base_path: Path):
    """Extract Dojo exam format - read from Word document to preserve formatting."""
    from docx import Document
    
    # Read from Word document instead of PDF to preserve formatting
    word_path = base_path / "pdf-formatting" / "word" / f"{filename}.docx"
    
    if not word_path.exists():
        raise FileNotFoundError(f"Word document not found: {word_path}")
    
    doc = Document(str(word_path))
    
    output_lines = []
    skip_content = False
    prev_was_list_item = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check if we hit "Explanations:" - start skipping
        if text == "Explanations:":
            skip_content = True
            continue
        
        # Check if we hit a new question - stop skipping
        if re.match(r'^Question\s+\d+:', text):
            skip_content = False
            # Add blank line before question if previous was a list item
            if prev_was_list_item:
                output_lines.append("\n")
                prev_was_list_item = False
        
        # Skip if we're in explanations section
        if skip_content:
            continue
        
        # Skip empty paragraphs
        if not text:
            continue
        
        # Check if paragraph is bold
        is_bold = any(run.bold for run in para.runs if run.text.strip())
        
        # Check if this is a list item
        if text.startswith('- '):
            if is_bold:
                # Bold list item - extract text after "- " and make it bold
                option_text = text[2:]
                output_lines.append(f"- **{option_text}**\n")
            else:
                output_lines.append(f"{text}\n")
            prev_was_list_item = True
        else:
            # Regular paragraph
            if is_bold:
                output_lines.append(f"**{text}**\n\n")
            else:
                output_lines.append(f"{text}\n\n")
            prev_was_list_item = False
    
    # Save to markdown file
    output_file = output_dir / f"{filename}.md"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(''.join(output_lines))
    
    print(f"Compact exam version saved to: {output_file}")
    
    # Generate PDF from the markdown
    pdf_file = generate_pdf_from_markdown_internal(output_file, filename, base_path)
    print(f"PDF version saved to: {pdf_file}")
    
    return output_file


def generate_pdf_from_markdown_internal(markdown_path: Path, filename: str, base_path: Path):
    """Convert Markdown compact exam to PDF (internal function).
    
    Args:
        markdown_path: Path object to the markdown file
        filename: Base filename without extension
        base_path: Base path for output
    """
    # Read markdown file
    with open(markdown_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content)
    
    # Add CSS styling
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                line-height: 1.6;
                margin: 40px;
            }}
            ul {{
                list-style-type: none;
                padding-left: 20px;
            }}
            li::before {{
                content: "• ";
                font-weight: bold;
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    # Generate PDF in compact-exam-versions folder
    pdf_output_dir = base_path / "pdf-formatting" / "compact-exam-versions"
    pdf_output_dir.mkdir(parents=True, exist_ok=True)
    output_pdf = pdf_output_dir / f"{filename}.pdf"
    HTML(string=styled_html).write_pdf(output_pdf)
    
    return str(output_pdf)


def generate_pdf_from_markdown(filename: str):
    """Convert Markdown compact exam to PDF.
    
    Args:
        filename: Name of the file without extension (e.g., 'exam')
    """
    # Get GNL_PROCESSING_PATH from environment
    gnl_processing_path = os.getenv('GNL_PROCESSING_PATH')
    if not gnl_processing_path:
        raise ValueError("GNL_PROCESSING_PATH not found in .env file")
    
    # Construct input path
    base_path = Path(gnl_processing_path).parent
    markdown_path = base_path / "pdf-formatting" / "compact-exam-versions" / f"{filename}.md"
    
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {markdown_path}")
    
    # Generate PDF
    output_pdf = generate_pdf_from_markdown_internal(markdown_path)
    
    print(f"PDF generated: {output_pdf}")
    return output_pdf


if __name__ == "__main__":
    fire.Fire(extract_compact_exam)
