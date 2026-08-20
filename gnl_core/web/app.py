"""GNL Web UI — FastAPI + HTMX dashboard with scheduler."""

import asyncio
import json
from contextlib import asynccontextmanager
from fastapi import FastAPI, Request, WebSocket, WebSocketDisconnect
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from pathlib import Path
import os
from dotenv import load_dotenv
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger

load_dotenv(Path(__file__).parent.parent.parent / '.env')

templates = Jinja2Templates(directory=str(Path(__file__).parent / "templates"))

ws_clients: list[WebSocket] = []
scheduler = AsyncIOScheduler()
_stop_signal = False
def _deliver_all_sync(loop=None):
    """Run deliver for all active parents (called by scheduler/startup)."""
    from gnl_core.db import get_active_parents, parent_status
    from gnl_core.generate import generate
    from gnl_core.download import download
    from gnl_core.convert import convert

    if loop is None:
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()

    def notify(msg):
        asyncio.run_coroutine_threadsafe(broadcast_log(msg), loop)
        asyncio.run_coroutine_threadsafe(broadcast_status(), loop)

    notify("🕐 Scheduled auto-deliver started")

    for pid in get_active_parents():
        s = parent_status(pid)
        if s['generated'] < s['total']:
            notify(f"▶ AUTO: GENERATE (parent {pid})")
            generate(pid, on_progress=lambda rec: asyncio.run_coroutine_threadsafe(broadcast_status(), loop))
            notify(f"✓ AUTO: Generate done (parent {pid})")
        s = parent_status(pid)
        if s['downloaded'] < s['generated']:
            notify(f"▶ AUTO: DOWNLOAD (parent {pid})")
            download(pid)
            notify(f"✓ AUTO: Download done (parent {pid})")
        s = parent_status(pid)
        if s['downloaded'] > 0 and s['converted'] < s['downloaded']:
            notify(f"▶ AUTO: CONVERT (parent {pid})")
            convert(pid)
            notify(f"✓ AUTO: Convert done (parent {pid})")


def _scheduled_linkedin_fetch():
    """Scheduled: fetch LinkedIn saved posts."""
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        asyncio.run_coroutine_threadsafe(_fetch_saved_articles('linkedin'), loop)
    except Exception:
        pass


def _scheduled_linkedin_generate():
    """Scheduled: batch generate for LinkedIn articles."""
    import asyncio
    try:
        loop = asyncio.get_event_loop()
        asyncio.run_coroutine_threadsafe(_batch_generate('linkedin'), loop)
    except Exception:
        pass


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Mount Google Drive if not available
    import subprocess
    if not os.path.ismount('/mnt/g'):
        subprocess.run(['sudo', 'mount', '-t', 'drvfs', 'G:', '/mnt/g'], capture_output=True)

    # Load scheduler config
    from gnl_core.config import get_config
    config = get_config()
    sched_config = config.get('SCHEDULER', {})
    if isinstance(sched_config, str):
        import json as _json
        sched_config = _json.loads(sched_config) if sched_config else {}

    # GNL Deliver
    gnl_job = sched_config.get('gnl_deliver', {})
    if gnl_job.get('enabled'):
        h, m = gnl_job.get('time', '08:00').split(':')
        scheduler.add_job(_deliver_all_sync, CronTrigger(hour=int(h), minute=int(m), timezone='America/Toronto'), id='gnl_deliver', replace_existing=True, misfire_grace_time=3600, args=[asyncio.get_event_loop()])

    # Saved Articles — Fetch
    fetch_job = sched_config.get('saved_articles_fetch', {})
    if fetch_job.get('enabled'):
        h, m = fetch_job.get('time', '02:00').split(':')
        scheduler.add_job(_scheduled_linkedin_fetch, CronTrigger(hour=int(h), minute=int(m), timezone='America/Toronto'), id='saved_articles_fetch', replace_existing=True, misfire_grace_time=3600)

    # Saved Articles — Generate
    gen_job = sched_config.get('saved_articles_generate', {})
    if gen_job.get('enabled'):
        h, m = gen_job.get('time', '03:00').split(':')
        scheduler.add_job(_scheduled_linkedin_generate, CronTrigger(hour=int(h), minute=int(m), timezone='America/Toronto'), id='saved_articles_generate', replace_existing=True, misfire_grace_time=3600)

    scheduler.start()

    # Auto-deliver at startup disabled — use scheduled time or manual button instead
    # from gnl_core.db import get_active_parents
    # if get_active_parents() and _get_quota() > 0:
    #     ...

    yield
    scheduler.shutdown()

app = FastAPI(title="GNL Process", lifespan=lifespan)


def _get_quota():
    from gnl_core.db import get_db
    from datetime import datetime, timezone as tz, timedelta
    max_quota = int(os.environ.get('TEST_QUOTA', '20')) if os.environ.get('TEST_MODE', '0') == '1' else 20
    pt = tz(timedelta(hours=-7))
    now_pt = datetime.now(pt)
    midnight_pt_str = now_pt.replace(hour=0, minute=0, second=0, microsecond=0).strftime('%Y-%m-%d')
    with get_db() as conn:
        used = conn.execute("SELECT COUNT(*) as c FROM podcast_download WHERE generation_state=1 AND date=?", (midnight_pt_str,)).fetchone()['c']
    return max(0, max_quota - used)
async def broadcast_log(msg: str):
    for ws in ws_clients[:]:
        try:
            await ws.send_text(json.dumps({"type": "log", "msg": msg}))
        except Exception:
            ws_clients.remove(ws)
async def broadcast_status():
    """Push updated editions HTML to all clients."""
    from gnl_core.db import get_db, parent_status

    quota_remaining = _get_quota()

    with get_db() as conn:
        rows = conn.execute("SELECT id, podcast_subtheme, parent_file FROM parent_configuration WHERE combination_state != -1").fetchall()

    html = ""
    for row in rows:
        s = parent_status(row['id'])
        # Skip completed editions (they go to history)
        if s['combined'] == 1 and s['converted'] == s['total'] and s['downloaded'] == s['total']:
            continue
        pid = row['id']
        sub = row['podcast_subtheme']
        total = s['total']
        gen = s['generated']
        dl = s['downloaded']
        conv = s['converted']

        # Button
        if s['combined'] == 1 and conv == total and dl == total:
            btn = '<span class="px-3 py-1 bg-green-900 text-green-300 rounded-full text-xs font-medium">✅ Terminé</span>'
        elif conv == total and total > 0:
            btn = f'<button hx-post="/action/deliver/{pid}" hx-swap="none" onclick="this.disabled=true;this.textContent=\'...\'" class="px-4 py-2 bg-green-600 hover:bg-green-500 rounded-lg text-sm font-medium">▶ Combiner</button>'
        elif gen == 0:
            btn = f'<button hx-post="/action/deliver/{pid}" hx-swap="none" onclick="this.disabled=true;this.textContent=\'...\'" class="px-4 py-2 bg-blue-600 hover:bg-blue-500 rounded-lg text-sm font-medium">▶ Lancer</button>'
        elif dl < total:
            btn = f'<button hx-post="/action/deliver/{pid}" hx-swap="none" onclick="this.disabled=true;this.textContent=\'...\'" class="px-4 py-2 bg-yellow-600 hover:bg-yellow-500 rounded-lg text-sm font-medium">▶ Reprendre</button>'
        else:
            btn = f'<button hx-post="/action/deliver/{pid}" hx-swap="none" onclick="this.disabled=true;this.textContent=\'...\'" class="px-4 py-2 bg-orange-600 hover:bg-orange-500 rounded-lg text-sm font-medium">▶ Finaliser</button>'

        clean_btn = f'<button hx-post="/action/clean/{pid}" hx-swap="none" onclick="this.disabled=true" class="px-2 py-1 text-red-400 hover:text-red-300 text-xs">🗑</button>'

        # Timeline steps
        steps = [("Préparé", total, total), ("Généré", gen, total), ("Téléchargé", dl, total), ("Converti", conv, total), ("Combiné", s['combined'], 1)]
        timeline = '<div class="flex items-center justify-between">'
        for i, (label, done, t) in enumerate(steps):
            if done == t:
                dot_cls = "bg-green-500 border-green-500"
                txt_cls = "text-green-400"
            elif done > 0:
                dot_cls = "bg-yellow-500 border-yellow-500 animate-pulse"
                txt_cls = "text-yellow-400"
            else:
                dot_cls = "bg-gray-700 border-gray-600"
                txt_cls = "text-gray-500"
            timeline += f'<div class="flex flex-col items-center z-10" style="width:20px"><div class="w-4 h-4 rounded-full border-2 {dot_cls}"></div><span class="text-xs mt-1 {txt_cls}">{label}</span><span class="text-xs text-gray-500">{done}/{t}</span></div>'
            if i < 4:
                next_done = steps[i+1][1]
                next_total = steps[i+1][2]
                pct = int(next_done / next_total * 100) if next_total > 0 else 0
                current_complete = (done == t)
                if current_complete:
                    bar_cls = "bg-green-500"
                    bar_pct = 100
                elif pct > 0:
                    bar_cls = "bg-yellow-500"
                    bar_pct = pct
                else:
                    bar_cls = ""
                    bar_pct = 0
                timeline += f'<div class="flex-1 h-1.5 -mt-6 bg-gray-700 rounded-full overflow-hidden mx-1"><div class="h-full rounded-full transition-all duration-500 {bar_cls}" style="width:{bar_pct}%"></div></div>'
        timeline += '</div>'

        quota_msg = f'<p class="text-xs text-gray-400 mt-3">⏳ Quota — reprend automatiquement demain</p>' if 0 < gen < total else ''

        html += f'''<div class="bg-gray-800 rounded-lg p-5">
            <div class="flex justify-between items-start mb-4">
                <div><h3 class="font-semibold text-lg">{row['parent_file']}</h3><span class="text-xs text-gray-400">{total} épisodes</span></div>
                <div class="flex items-center gap-2">{btn}{clean_btn}</div>
            </div>
            <div class="relative">{timeline}</div>{quota_msg}
        </div>'''

    if not html:
        html = '<div class="bg-gray-800 rounded-lg p-8 text-center text-gray-500">Aucune édition.</div>'

    # Build history HTML
    with get_db() as conn:
        all_rows = conn.execute("SELECT id, podcast_subtheme, parent_file FROM parent_configuration").fetchall()
    history_html = ""
    for row in all_rows:
        s = parent_status(row['id'])
        if s['combined'] == -1:
            history_html += f'<div class="bg-gray-800 rounded-lg p-4 flex justify-between items-center"><div><h3 class="font-semibold">{row["parent_file"]}</h3><span class="text-xs text-gray-400">{s["total"]} épisodes · {row["podcast_subtheme"]}</span></div><span class="px-3 py-1 bg-red-900 text-red-300 rounded-full text-xs">🗑 Supprimé</span></div>'
        elif s['combined'] == 1 and s['converted'] == s['total'] and s['downloaded'] == s['total']:
            history_html += f'<div class="bg-gray-800 rounded-lg p-4 flex justify-between items-center"><div><h3 class="font-semibold">{row["parent_file"]}</h3><span class="text-xs text-gray-400">{s["total"]} épisodes · {row["podcast_subtheme"]}</span></div><span class="px-3 py-1 bg-green-900 text-green-300 rounded-full text-xs">✅ Terminé</span></div>'
    if not history_html:
        history_html = '<div class="bg-gray-800 rounded-lg p-8 text-center text-gray-500">Aucune édition terminée.</div>'

    for ws in ws_clients[:]:
        try:
            await ws.send_text(json.dumps({"type": "status_update", "html": html, "history_html": history_html, "quota": quota_remaining}))
        except Exception:
            ws_clients.remove(ws)
@app.get("/", response_class=HTMLResponse)
async def dashboard(request: Request):
    from gnl_core.db import get_db, parent_status
    from datetime import datetime, timezone, timedelta

    with get_db() as conn:
        rows = conn.execute("SELECT id, podcast_subtheme, parent_file FROM parent_configuration").fetchall()

    parents = []
    history = []
    for row in rows:
        s = parent_status(row['id'])
        item = {'id': row['id'], 'subtheme': row['podcast_subtheme'], 'parent_file': row['parent_file'], **s}
        if s['combined'] == -1:
            item['status_tag'] = 'deleted'
            history.append(item)
        elif s['combined'] == 1 and s['converted'] == s['total'] and s['downloaded'] == s['total']:
            item['status_tag'] = 'completed'
            history.append(item)
        else:
            parents.append(item)

    schedule_time = os.getenv('GNL_SCHEDULE_TIME', '08:00')
    jobs = scheduler.get_jobs()
    next_run = str(jobs[0].next_run_time.strftime('%Y-%m-%d %H:%M')) if jobs else "Not scheduled"
    test_mode = os.environ.get('TEST_MODE', '0') == '1'

    quota_remaining = _get_quota()

    # Config for admin tab
    from gnl_core.config import get_config
    config = get_config()

    # Changelog
    import markdown
    changelog_path = Path(__file__).parent.parent.parent / 'CHANGELOG.md'
    changelog_html = markdown.markdown(changelog_path.read_text()) if changelog_path.exists() else "<p>No changelog found.</p>"

    return templates.TemplateResponse("dashboard.html", {
        "request": request, "parents": parents, "history": history,
        "schedule_time": schedule_time, "next_run": next_run, "test_mode": test_mode,
        "quota_remaining": quota_remaining, "config": config, "changelog_html": changelog_html
    })
@app.get("/api/catalog")
async def get_catalog():
    """Return series catalog for form dropdowns."""
    from gnl_core.db import get_db
    with get_db() as conn:
        rows = conn.execute("SELECT theme, subtheme, content_mode FROM series_catalog ORDER BY theme, id").fetchall()
    catalog = {}
    for r in rows:
        catalog.setdefault(r['theme'], []).append({'subtheme': r['subtheme'], 'mode': r['content_mode'] or 'manual'})
    return catalog


@app.get("/api/prompts")
async def get_prompts():
    """Return all prompts."""
    from gnl_core.db import get_db
    with get_db() as conn:
        rows = conn.execute("SELECT id, theme, subtheme, prompt FROM series_catalog ORDER BY theme, id").fetchall()
    return [dict(r) for r in rows]


@app.post("/api/catalog/add")
async def add_catalog_entry(request: Request):
    """Add a theme/subtheme to the catalog."""
    from gnl_core.db import get_db
    form = await request.form()
    theme = form.get("theme", "").strip()
    subtheme = form.get("subtheme", "").strip()
    if not theme or not subtheme:
        return {"status": "error", "message": "Theme and subtheme required"}
    with get_db() as conn:
        conn.execute("INSERT OR IGNORE INTO series_catalog (theme, subtheme, prompt) VALUES (?, ?, '')", (theme, subtheme))
        conn.commit()
    await broadcast_log(f"✓ Ajouté: {theme}/{subtheme}")
    return {"status": "ok"}


@app.post("/api/catalog/delete/{catalog_id}")
async def delete_catalog_entry(catalog_id: int):
    """Delete a catalog entry."""
    from gnl_core.db import get_db
    with get_db() as conn:
        conn.execute("DELETE FROM series_catalog WHERE id=?", (catalog_id,))
        conn.commit()
    await broadcast_log("✓ Sous-thème supprimé")
    return {"status": "ok"}


@app.post("/api/prompts/{catalog_id}")
async def save_prompt(catalog_id: int, request: Request):
    """Save prompt for a catalog entry."""
    from gnl_core.db import get_db
    form = await request.form()
    prompt = form.get("prompt", "")
    with get_db() as conn:
        conn.execute("UPDATE series_catalog SET prompt=? WHERE id=?", (prompt, catalog_id))
        conn.commit()
    await broadcast_log("✓ Prompt sauvegardé")
    return {"status": "ok"}
@app.post("/stop")
async def stop_processing():
    """Signal all running operations to stop."""
    global _stop_signal
    _stop_signal = True
    await broadcast_log("⏹ Stop signal sent")
    return {"status": "stopping"}


@app.get("/content/files/{theme}/{subtheme}")
async def list_content_files(theme: str, subtheme: str):
    """List PDFs in INBOX_FOLDER/{theme}/{subtheme}/ with processed status."""
    from gnl_core.config import get_config
    from gnl_core.db import get_db
    config = get_config()
    inbox = config.get('INBOX_FOLDER', '')
    folder = os.path.join(inbox, theme, subtheme)
    if not os.path.isdir(folder):
        return []

    backlog = config.get('GNL_BACKLOG', '')
    backlog_dir = os.path.join(backlog, theme, subtheme)
    delivered = set()
    if os.path.isdir(backlog_dir):
        delivered = {os.path.splitext(f)[0] for f in os.listdir(backlog_dir) if f.lower().endswith('.mp3')}

    files = []
    for f in sorted(os.listdir(folder)):
        if f.lower().endswith(('.pdf', '.docx')) and not f.startswith('~$'):
            name_no_ext = os.path.splitext(f)[0]
            item = {'name': f, 'processed': name_no_ext in delivered}
            
            if theme == 'exams':
                # Check if Anki was generated
                from gnl_core.exams import get_exam_base
                base = get_exam_base(theme, subtheme)
                anki_path = base / 'Anki-generation' / 'anki' / f"{name_no_ext}.apkg"
                item['anki'] = anki_path.exists()
                # Check if in production (exists in parent_configuration)
                with get_db() as conn:
                    in_prod = conn.execute("SELECT id FROM parent_configuration WHERE parent_file=? AND podcast_subtheme=?", (name_no_ext, subtheme)).fetchone()
                item['in_production'] = in_prod is not None
            
            files.append(item)
    return files


@app.post("/exams/clean/{theme}/{subtheme}/{filename}")
async def clean_exam(theme: str, subtheme: str, filename: str):
    """Remove exam file and all its generated variants."""
    from gnl_core.config import get_config
    from gnl_core.exams import get_exam_base

    config = get_config()
    name = os.path.splitext(filename)[0]
    base = get_exam_base(theme, subtheme)

    removed = 0
    # Only remove generated variants in assets/ — keep original in inbox
    # origin/
    for ext in ['.docx', '.pdf']:
        p = base / 'pdf-formatting' / 'origin' / f"{name}{ext}"
        if p.exists():
            p.unlink()
            removed += 1
    # word/
    p = base / 'pdf-formatting' / 'word' / f"{name}.docx"
    if p.exists():
        p.unlink()
        removed += 1
    # full-markdown/
    p = base / 'pdf-formatting' / 'full-markdown' / f"{name}.md"
    if p.exists():
        p.unlink()
        removed += 1
    # Anki-generation/markdown/
    p = base / 'Anki-generation' / 'markdown' / f"{name}.md"
    if p.exists():
        p.unlink()
        removed += 1
    # Anki-generation/anki/
    for ext in ['-anki.txt', '.apkg']:
        p = base / 'Anki-generation' / 'anki' / f"{name}{ext}"
        if p.exists():
            p.unlink()
            removed += 1

    # PDF_PARTS_FOLDER/{theme}/{subtheme}/{name}/
    import shutil
    pdf_parts = config.get('PDF_PARTS_FOLDER', '')
    if pdf_parts:
        pdf_dir = os.path.join(pdf_parts, theme, subtheme, name)
        if os.path.isdir(pdf_dir):
            shutil.rmtree(pdf_dir)
            removed += 1

    # AUDIO_PARTS_FOLDER/{theme}/{subtheme}/{name}/
    audio_parts = config.get('AUDIO_PARTS_FOLDER', '')
    if audio_parts:
        audio_dir = os.path.join(audio_parts, theme, subtheme, name)
        if os.path.isdir(audio_dir):
            shutil.rmtree(audio_dir)
            removed += 1

    # Delete NLM notebook ({name}-FULL)
    try:
        from notebooklm_tools.mcp.tools._utils import get_client
        from notebooklm_tools.services.notebooks import list_notebooks, delete_notebook
        client = get_client()
        notebook_title = f"{name}-FULL"
        nbs = list_notebooks(client)
        for nb in nbs.get('notebooks', []):
            if nb.get('title') == notebook_title:
                delete_notebook(client, nb['id'])
                removed += 1
                break
    except Exception:
        pass

    # Remove from DB (parent_configuration + podcast_download)
    from gnl_core.db import get_db
    with get_db() as conn:
        row = conn.execute("SELECT id FROM parent_configuration WHERE parent_file=? AND podcast_subtheme=?", (name, subtheme)).fetchone()
        if row:
            conn.execute("DELETE FROM podcast_download WHERE parent_configuration_id=?", (row['id'],))
            conn.execute("DELETE FROM parent_configuration WHERE id=?", (row['id'],))
            conn.commit()
            removed += 1

    await broadcast_log(f"🗑 Nettoyé: {name} ({removed} fichiers/dossiers supprimés)")
    return {"status": "ok", "removed": removed}


@app.post("/exams/process/{theme}/{subtheme}/{filename}")
async def process_exam(theme: str, subtheme: str, filename: str):
    """Run full exam pipeline: format → highlight → compact → anki."""
    asyncio.create_task(_process_exam(theme, subtheme, filename))
    return {"status": "started"}


async def _process_exam(theme, subtheme, filename):
    from gnl_core.config import get_config
    config = get_config()
    inbox = config.get('INBOX_FOLDER', '')
    file_path = os.path.join(inbox, theme, subtheme, filename)

    if not os.path.isfile(file_path):
        await broadcast_log(f"⚠ Fichier introuvable: {file_path}")
        await broadcast_log("__done__")
        return

    loop = asyncio.get_event_loop()
    name = os.path.splitext(filename)[0]
    origin = 'dojo' if 'dojo' in filename.lower() else 'udemy'

    try:
        from gnl_core.exams import get_exam_base, step1_format, step2b_full_markdown, step3_highlight, step5_anki
        import shutil

        # Move file to assets/pdf-formatting/origin/
        base = get_exam_base(theme, subtheme)
        origin_dir = base / 'pdf-formatting' / 'origin'
        origin_dir.mkdir(parents=True, exist_ok=True)
        origin_path = origin_dir / filename

        if not origin_path.exists():
            shutil.copy2(file_path, str(origin_path))
            await broadcast_log(f"📂 Copié vers origin/{filename}")

        # Step 1: origin/ → word/
        await broadcast_log(f"▶ [1/5] FORMAT (origin → word)")

        def on_p1(msg):
            asyncio.run_coroutine_threadsafe(broadcast_log(f"  {msg}"), loop)

        word_path = await loop.run_in_executor(None, lambda: step1_format(str(origin_path), theme, subtheme, origin, on_progress=on_p1))
        await broadcast_log(f"  ✓ {word_path}")

        # Step 2: word/ → markdown/
        await broadcast_log("▶ [2/5] MARKDOWN (word → md)")
        md_path = await loop.run_in_executor(None, lambda: step2b_full_markdown(word_path, theme, subtheme, on_progress=on_p1))
        await broadcast_log(f"  ✓ {md_path}")

        # Step 3: highlight (Bedrock)
        await broadcast_log("▶ [3/5] HIGHLIGHT (Bedrock → correct answers)")
        answers = await loop.run_in_executor(None, lambda: step3_highlight(md_path, on_progress=on_p1))
        await broadcast_log(f"  ✓ {len(answers)} questions analysées")

        # Step 4: anki cards (directly from answers)
        await broadcast_log("▶ [4/4] ANKI (flashcards)")
        anki_path = await loop.run_in_executor(None, lambda: step5_anki(answers, md_path, theme, subtheme, on_progress=on_p1))
        await broadcast_log(f"  ✓ {anki_path}")

        await broadcast_log(f"✓ Pipeline terminé: {name}")
    except Exception as e:
        await broadcast_log(f"⚠ Erreur: {str(e)[:100]}")
    await broadcast_log("__done__")


@app.post("/content/interactive")
async def launch_interactive(request: Request):
    """Create a notebook with full PDF and generate audio in English for interactive mode."""
    form = await request.form()
    theme = form.get("theme", "")
    subtheme = form.get("subtheme", "")
    filename = form.get("filename", "")

    asyncio.create_task(_launch_interactive(theme, subtheme, filename))
    return {"status": "started"}


async def _launch_interactive(theme, subtheme, filename):
    from gnl_core.config import get_config
    config = get_config()
    inbox = config.get('INBOX_FOLDER', '')
    pdf_path = os.path.join(inbox, theme, subtheme, filename)

    if not os.path.isfile(pdf_path):
        await broadcast_log(f"⚠ Fichier introuvable: {pdf_path}")
        await broadcast_log("__done__")
        return

    loop = asyncio.get_event_loop()
    name = os.path.splitext(filename)[0]
    await broadcast_log(f"▶ Mode interactif: {filename} (anglais, vitesse 1x)")

    try:
        def _do():
            from notebooklm_tools.mcp.tools._utils import get_client
            from notebooklm_tools.services.notebooks import create_notebook, list_notebooks
            from notebooklm_tools.services.sources import add_source
            from notebooklm_tools.services.studio import create_artifact

            client = get_client()
            nb = create_notebook(client, f"[Interactive] {name}")
            nb_id = nb['notebook_id']
            add_source(client, nb_id, "file", file_path=pdf_path, wait=True)
            create_artifact(client, nb_id, "audio", language="en")
            return nb.get('url', f"https://notebooklm.google.com/notebook/{nb_id}")

        url = await loop.run_in_executor(None, _do)
        await broadcast_log(f"✓ Notebook prêt: {url}")
        await broadcast_log(f"🔗 <a href='{url}' target='_blank' class='text-blue-400 underline'>{url}</a>")
    except Exception as e:
        await broadcast_log(f"⚠ Erreur: {str(e)[:100]}")
    await broadcast_log("__done__")


@app.get("/saved-articles/{source}")
async def get_saved_articles(source: str):
    """List saved articles for a source (linkedin, medium, etc.)."""
    from gnl_core.db import get_db
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, title, source_url, saved_date, processed, output_path FROM saved_articles WHERE source=? ORDER BY saved_date DESC, id ASC",
            (source,)
        ).fetchall()
    return [dict(r) for r in rows]


@app.post("/saved-articles/fetch/{source}")
async def fetch_saved_articles(source: str):
    """Fetch saved items from source (linkedin, medium) via MCP."""
    asyncio.create_task(_fetch_saved_articles(source))
    return {"status": "started"}


async def _fetch_saved_articles(source):
    await broadcast_log(f"▶ Fetch saved articles ({source})...")

    if source == 'linkedin':
        loop = asyncio.get_event_loop()
        # Step 1: Call LinkedIn MCP to refresh cache
        await broadcast_log("🔄 Appel MCP LinkedIn (scraping)...")
        refresh_ok = await _call_linkedin_mcp()
        if refresh_ok:
            await broadcast_log("✓ Cache LinkedIn mis à jour")
        else:
            await broadcast_log("⚠ Scraping échoué — utilisation du cache existant")

        # Step 2: Import from cache to our DB
        count = await loop.run_in_executor(None, _fetch_linkedin_from_cache)
        if count >= 0:
            await broadcast_log(f"✓ {count} nouveaux articles importés")
        else:
            await broadcast_log("⚠ Cache LinkedIn introuvable")
    else:
        await broadcast_log(f"⚠ Fetch {source}: pas encore implémenté")
    await broadcast_log("__done__")


async def _call_linkedin_mcp():
    """Call LinkedIn MCP server to refresh saved posts cache."""
    try:
        from mcp import ClientSession, StdioServerParameters
        from mcp.client.stdio import stdio_client
        from pathlib import Path

        # Delete old cache to get a fresh ordered scrape
        cache_db = Path.home() / ".linkedin-mcp" / "saved_posts.db"
        if cache_db.exists():
            cache_db.unlink()

        server_params = StdioServerParameters(
            command='python3',
            args=['-m', 'linkedin_mcp_server'],
            env={**os.environ, 'PYTHONPATH': os.environ.get('LINKEDIN_MCP_PATH', '/home/nizar/HomeWspce/linkedin-mcp-fork')}
        )

        async with stdio_client(server_params) as (read, write):
            async with ClientSession(read, write) as session:
                await session.initialize()
                await session.call_tool('get_saved_posts', {'num_posts': int(os.environ.get('LINKEDIN_SCRAPE_COUNT', '50')), 'full_content': True})
        return True
    except Exception as e:
        return False


def _generate_title_bedrock(content: str) -> str:
    """Generate a meaningful title for an article using Bedrock (Claude)."""
    try:
        import boto3, json
        from gnl_core.config import get_config
        config = get_config()
        model_id = config.get('BEDROCK_MODEL_ID', 'us.anthropic.claude-sonnet-4-20250514-v1:0')
        client = boto3.client('bedrock-runtime', region_name='us-east-1')
        # Use first 800 chars for context
        snippet = content[:800]
        response = client.invoke_model(
            modelId=model_id,
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 50,
                "messages": [{"role": "user", "content": f"Give a concise title (max 10 words, no quotes) that summarizes the main topic of this LinkedIn post:\n\n{snippet}"}]
            })
        )
        result = json.loads(response['body'].read())
        title = result['content'][0]['text'].strip().strip('"\'')
        return title[:100] if title else 'Sans titre'
    except Exception:
        return ''


def _fetch_linkedin_from_cache():
    """Read LinkedIn MCP cache and import into our saved_articles table."""
    from gnl_core.db import get_db
    from pathlib import Path
    import sqlite3 as sqlite
    from datetime import datetime, timedelta

    cache_db = Path.home() / ".linkedin-mcp" / "saved_posts.db"
    if not cache_db.exists():
        return -1

    # Read from LinkedIn cache (ordered by id ASC = top of page = most recent)
    conn_cache = sqlite.connect(cache_db)
    conn_cache.row_factory = sqlite.Row
    rows = conn_cache.execute("SELECT author, content, url, scraped_at FROM saved_posts ORDER BY id ASC").fetchall()
    conn_cache.close()

    # Import into our DB with position (for ordering)
    added = 0
    now = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    today = datetime.now()
    with get_db() as conn:
        for position, r in enumerate(rows):
            url = r['url'] or ''
            if not url or url.startswith('no-url'):
                continue
            content = r['content'] or ''
            # Generate title via Bedrock (fallback to first substantial line)
            title = _generate_title_bedrock(content)
            if not title:
                lines = content.split('\n')
                title = 'Sans titre'
                for l in lines:
                    l = l.strip()
                    if len(l) > 40 and '•' not in l and 'abonnés' not in l and 'Architect' not in l and 'Engineer' not in l and 'Creator' not in l:
                        title = l[:100]
                        break
            # Extract relative date from content (e.g. "• 1 j", "• 2 sem.", "• 3 h")
            import re
            date_str = ''
            date_match = re.search(r'•\s*(\d+)\s*(j|h|sem|mois|min)', content)
            if date_match:
                num = int(date_match.group(1))
                unit = date_match.group(2)
                if unit == 'h' or unit == 'min':
                    post_date = today
                elif unit == 'j':
                    post_date = today - timedelta(days=num)
                elif unit == 'sem':
                    post_date = today - timedelta(weeks=num)
                elif unit == 'mois':
                    post_date = today - timedelta(days=num * 30)
                else:
                    post_date = today
                date_str = post_date.strftime('%Y-%m-%d')
            else:
                date_str = today.strftime('%Y-%m-%d')
            try:
                conn.execute(
                    "INSERT OR IGNORE INTO saved_articles (source, source_id, title, content, source_url, saved_date, fetched_at, processed) VALUES (?, ?, ?, ?, ?, ?, ?, 0)",
                    ('linkedin', url, title, content, url, date_str, now)
                )
                # Update title for existing articles if Bedrock generated a better one
                if title and title != 'Sans titre':
                    conn.execute(
                        "UPDATE saved_articles SET title = ? WHERE source_id = ? AND (title = 'Sans titre' OR length(title) > 60)",
                        (title, url)
                    )
                added += 1
                # Save original content to INBOX_FOLDER
                from gnl_core.config import get_config
                config = get_config()
                inbox = config.get('INBOX_FOLDER', '')
                raw_dir = os.path.join(inbox, 'saved-articles', 'linkedin')
                os.makedirs(raw_dir, exist_ok=True)
                safe = "".join(c if c.isalnum() or c in '-_ ' else '' for c in title)[:60]
                raw_path = os.path.join(raw_dir, f"{safe}.txt")
                if not os.path.exists(raw_path):
                    with open(raw_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            except Exception:
                pass
        conn.commit()
    return added


@app.get("/saved-articles/audio-files/{source}")
async def list_audio_files(source: str):
    """List MP3 files not yet combined, sorted by creation date."""
    from gnl_core.db import get_db
    from gnl_core.config import get_config
    config = get_config()
    audio_dir = os.path.join(config.get('AUDIO_PARTS_FOLDER', ''), 'saved-articles', source)

    # Get uncombined articles from DB
    with get_db() as conn:
        rows = conn.execute(
            "SELECT audio_path FROM saved_articles WHERE source=? AND processed=1 AND audio_path IS NOT NULL AND combined_at IS NULL",
            (source,)
        ).fetchall()
    uncombined = {r[0] for r in rows if r[0]}

    if not os.path.isdir(audio_dir):
        return []
    files = []
    for f in os.listdir(audio_dir):
        if f.endswith('.mp3'):
            path = os.path.join(audio_dir, f)
            # Only show if not already combined
            if path not in uncombined:
                continue
            ctime = os.path.getctime(path)
            from datetime import datetime
            files.append({'filename': f, 'date': datetime.fromtimestamp(ctime).strftime('%Y-%m-%d %H:%M')})
    files.sort(key=lambda x: x['date'], reverse=True)
    return files


@app.post("/saved-articles/combine/{source}")
async def combine_articles(source: str, request: Request):
    """Combine selected MP3 files into a batch."""
    data = await request.json()
    selected_files = data.get('files', [])
    asyncio.create_task(_combine_selected(source, selected_files))
    return {"status": "combining"}


async def _combine_selected(source: str, selected_files: list):
    """Combine selected MP3 files into one batch file on the backlog drive."""
    import subprocess
    from gnl_core.config import get_config
    from datetime import datetime
    config = get_config()
    audio_dir = os.path.join(config.get('AUDIO_PARTS_FOLDER', ''), 'saved-articles', source)
    backlog_dir = os.path.join(config.get('GNL_BACKLOG', ''), 'saved-articles', source)
    os.makedirs(backlog_dir, exist_ok=True)

    # Build list of full paths (in selection order)
    paths = []
    for f in selected_files:
        p = os.path.join(audio_dir, f)
        if os.path.exists(p):
            paths.append(p)

    if not paths:
        await broadcast_log("⚠ Aucun fichier valide sélectionné")
        await broadcast_log("__done__")
        return

    await broadcast_log(f"▶ COMBINE ({len(paths)} fichiers)")

    date_str = datetime.now().strftime('%Y-%m-%d')
    output_file = os.path.join(backlog_dir, f"batch-{date_str}.mp3")
    counter = 1
    while os.path.exists(output_file):
        counter += 1
        output_file = os.path.join(backlog_dir, f"batch-{date_str}-{counter}.mp3")

    # Create ffmpeg concat file with silence between articles
    import tempfile
    silence_file = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'assets', 'silence-3s.mp3')
    concat_path = os.path.join(tempfile.gettempdir(), 'combine_concat.txt')
    with open(concat_path, 'w') as f:
        for idx, p in enumerate(paths):
            f.write(f"file '{p}'\n")
            if idx < len(paths) - 1 and os.path.exists(silence_file):
                f.write(f"file '{silence_file}'\n")

    result = subprocess.run(['ffmpeg', '-y', '-f', 'concat', '-safe', '0', '-i', concat_path, '-c', 'copy', output_file], capture_output=True)
    os.unlink(concat_path)

    if result.returncode == 0:
        # Mark combined in DB
        from gnl_core.db import get_db
        from datetime import datetime
        now_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S')
        with get_db() as conn:
            for p in paths:
                conn.execute("UPDATE saved_articles SET combined_at=? WHERE audio_path=?", (now_str, p))
            conn.commit()
        await broadcast_log(f"✓ Batch combiné: {output_file}")
    else:
        await broadcast_log(f"⚠ Erreur ffmpeg: {result.stderr.decode()[:100]}")
    await broadcast_log("__done__")


@app.post("/saved-articles/batch/{source}")
async def batch_generate_articles(source: str):
    """Generate text + audio for next 10 unprocessed articles."""
    asyncio.create_task(_batch_generate(source))
    return {"status": "started"}


async def _batch_generate(source):
    from gnl_core.db import get_db
    from gnl_core.config import get_config
    import subprocess

    with get_db() as conn:
        batch_size = int(os.environ.get('LINKEDIN_BATCH_SIZE', '10'))
        rows = conn.execute(
            "SELECT id, title, content FROM saved_articles WHERE source=? AND processed=0 ORDER BY saved_date DESC LIMIT ?",
            (source, batch_size)
        ).fetchall()

    if not rows:
        await broadcast_log("⚠ Aucun article à traiter")
        await broadcast_log("__done__")
        return

    total = len(rows)
    await broadcast_log(f"▶ Batch: {total} articles à traiter")

    generated_audio_paths = []

    for i, row in enumerate(rows):
        article_id = row['id']
        article_title = row['title'] or ''
        article_content = row['content'] or ''

        await broadcast_log(f"▶ [{i+1}/{total}] {article_title[:50]}...")

        # Step 1: Generate tunisian explanation
        loop = asyncio.get_event_loop()
        try:
            import boto3
            from botocore.config import Config as BotoConfig

            def _gen_text():
                from gnl_core.config import get_config
                config = get_config()
                model_id = config.get('BEDROCK_MODEL_ID', 'us.anthropic.claude-sonnet-4-20250514-v1:0')
                region = config.get('AWS_REGION', 'ca-central-1')
                profile = config.get('AWS_PROFILE', '')

                session = boto3.Session(profile_name=profile, region_name=region)
                client = session.client('bedrock-runtime', config=BotoConfig(read_timeout=600))

                prompt = f"""أنت خبير تقني تشرح بالعربي الدارج التونسي.

القواعد الصارمة:
- اكتب بالحروف العربية فقط (مش بالحروف اللاتينية/الفرنسية)
- المصطلحات التقنية بالإنجليزية كيما هي (API, cloud, container, deployment, agent, model...)
- الأسلوب: نثر محادثة طبيعي، كأنك تشرح لزميلك التونسي شفاهياً
- كل فقرة = جملة كاملة بالتونسي، والمصطلح الإنجليزي يتحط في وسط الجملة بشكل طبيعي
- اشرح كل النقاط بالتفصيل (مش ملخص)
- ما تترجمش المصطلحات التقنية — خليها بالإنجليزية
- كل مفهوم، كل أداة، كل ممارسة لازم تتشرح

❌ ممنوع:
- الكتابة بالحروف اللاتينية/الفرنسية
- الفصحى الكلاسيكية
- قوائم بنمط [مصطلح]: [شرح]
- ترجمة المصطلحات التقنية

المقال باش تشرحو:

العنوان: {article_title}
المحتوى: {article_content}"""

                response = client.converse(
                    modelId=model_id,
                    messages=[{"role": "user", "content": [{"text": prompt}]}],
                    inferenceConfig={"maxTokens": int(os.environ.get('BEDROCK_MAX_TOKENS', '4096'))}
                )
                return response['output']['message']['content'][0]['text']

            text_result = await loop.run_in_executor(None, _gen_text)

            # Save explanation to PDF_PARTS_FOLDER
            config = get_config()
            text_dir = os.path.join(config.get('PDF_PARTS_FOLDER', ''), 'saved-articles', source)
            os.makedirs(text_dir, exist_ok=True)
            safe_title = "".join(c if c.isalnum() or c in '-_ ' else '' for c in article_title)[:60]
            text_path = os.path.join(text_dir, f"{safe_title}.txt")
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(text_result)

        except Exception as e:
            await broadcast_log(f"  ⚠ Texte échoué: {str(e)[:60]}")
            continue

        # Step 2: TTS
        try:
            import requests
            import base64

            api_key = os.environ.get('GOOGLE_AI_API_KEY', '')
            if not api_key:
                from dotenv import load_dotenv
                load_dotenv()
                api_key = os.environ.get('GOOGLE_AI_API_KEY', '')

            tts_model = os.environ.get('TTS_MODEL', 'gemini-2.5-flash-preview-tts')
            url = f'https://generativelanguage.googleapis.com/v1beta/models/{tts_model}:generateContent?key={api_key}'
            payload = {
                'contents': [{'parts': [{'text': text_result}]}],
                'generationConfig': {
                    'responseModalities': ['AUDIO'],
                    'speechConfig': {
                        'voiceConfig': {
                            'prebuiltVoiceConfig': {'voiceName': os.environ.get('TTS_VOICE', 'Orus')}
                        }
                    }
                }
            }

            def _call_tts():
                resp = requests.post(url, json=payload, timeout=300)
                if resp.status_code != 200:
                    return None
                data = resp.json()
                audio_b64 = data['candidates'][0]['content']['parts'][0]['inlineData']['data']
                return base64.b64decode(audio_b64)

            audio_bytes = await loop.run_in_executor(None, _call_tts)
            if not audio_bytes:
                await broadcast_log(f"  ⚠ Audio échoué")
                continue

            # Save MP3
            audio_dir = os.path.join(config.get('AUDIO_PARTS_FOLDER', ''), 'saved-articles', source)
            os.makedirs(audio_dir, exist_ok=True)
            audio_path = os.path.join(audio_dir, f"{safe_title}.mp3")

            pcm_path = audio_path.replace('.mp3', '.pcm')
            with open(pcm_path, 'wb') as f:
                f.write(audio_bytes)
            subprocess.run(['ffmpeg', '-y', '-f', 's16le', '-ar', '24000', '-ac', '1', '-i', pcm_path, audio_path], capture_output=True)
            os.unlink(pcm_path)

            generated_audio_paths.append(audio_path)

            # Mark processed
            with get_db() as conn:
                conn.execute("UPDATE saved_articles SET processed=1, output_path=?, audio_path=? WHERE id=?", (text_path, audio_path, article_id))
                conn.commit()

            await broadcast_log(f"  ✓ [{i+1}/{total}] OK")
            await broadcast_status()

            # Rate limit delay between TTS calls
            tts_delay = int(os.environ.get('TTS_DELAY_SECONDS', '15'))
            if i < total - 1 and tts_delay > 0:
                await asyncio.sleep(tts_delay)

        except Exception as e:
            await broadcast_log(f"  ⚠ Audio échoué: {str(e)[:60]}")
            continue

    await broadcast_log("__done__")


@app.post("/saved-articles/generate/{article_id}")
async def generate_article_explanation(article_id: int):
    """Generate tunisian explanation for a saved article."""
    asyncio.create_task(_generate_article(article_id))
    return {"status": "started"}


async def _generate_article(article_id: int):
    from gnl_core.db import get_db
    import boto3

    with get_db() as conn:
        row = conn.execute("SELECT * FROM saved_articles WHERE id=?", (article_id,)).fetchone()
    if not row:
        await broadcast_log(f"⚠ Article {article_id} non trouvé")
        await broadcast_log("__done__")
        return

    await broadcast_log(f"▶ Génération explication tunisienne: {row['title'][:50]}...")
    loop = asyncio.get_event_loop()

    # Extract values before passing to thread (Row object not thread-safe)
    article_title = row['title'] or ''
    article_content = row['content'] or ''

    try:
        def _do():
            from gnl_core.config import get_config
            config = get_config()
            model_id = config.get('BEDROCK_MODEL_ID', 'us.anthropic.claude-sonnet-4-20250514-v1:0')
            region = config.get('AWS_REGION', 'ca-central-1')
            profile = config.get('AWS_PROFILE', '')

            session = boto3.Session(profile_name=profile, region_name=region)
            from botocore.config import Config
            client = session.client('bedrock-runtime', config=Config(read_timeout=600))

            prompt = f"""أنت خبير تقني تشرح بالعربي الدارج التونسي.

القواعد الصارمة:
- اكتب بالحروف العربية فقط (مش بالحروف اللاتينية/الفرنسية)
- المصطلحات التقنية بالإنجليزية كيما هي (API, cloud, container, deployment, agent, model...)
- الأسلوب: نثر محادثة طبيعي، كأنك تشرح لزميلك التونسي شفاهياً
- كل فقرة = جملة كاملة بالتونسي، والمصطلح الإنجليزي يتحط في وسط الجملة بشكل طبيعي
- اشرح كل النقاط بالتفصيل (مش ملخص)
- ما تترجمش المصطلحات التقنية — خليها بالإنجليزية
- كل مفهوم، كل أداة، كل ممارسة لازم تتشرح

❌ ممنوع:
- الكتابة بالحروف اللاتينية/الفرنسية (مثلاً: "hedhi la phase elli...")
- الفصحى الكلاسيكية
- قوائم بنمط [مصطلح]: [شرح]
- ترجمة المصطلحات التقنية

✅ مثال صحيح:
في الجزء متاع ال Monitoring، ما ضفناش alerting على ال FailedInvocations ولا على ال ApproximateAgeOfOldestMessage متاع ال bus. الفكرة أنو ال service هذي تخدم كـ proxy بين ال clients وال backend services.

المقال باش تشرحو:

العنوان: {article_title}
المحتوى: {article_content}"""

            response = client.converse(
                modelId=model_id,
                messages=[{"role": "user", "content": [{"text": prompt}]}],
                inferenceConfig={"maxTokens": int(os.environ.get('BEDROCK_MAX_TOKENS', '4096'))}
            )
            return response['output']['message']['content'][0]['text']

        result = await loop.run_in_executor(None, _do)

        # Save output (tunisian explanation → PDF_PARTS_FOLDER)
        from gnl_core.config import get_config
        config = get_config()
        output_dir = os.path.join(config.get('PDF_PARTS_FOLDER', ''), 'saved-articles', 'linkedin')
        os.makedirs(output_dir, exist_ok=True)
        safe_title = "".join(c if c.isalnum() or c in '-_ ' else '' for c in (article_title or f'article-{article_id}'))[:60]
        output_path = os.path.join(output_dir, f"{safe_title}.txt")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result)

        with get_db() as conn:
            conn.execute("UPDATE saved_articles SET processed=1, output_path=? WHERE id=?", (output_path, article_id))
            conn.commit()

        await broadcast_log(f"✓ Texte généré: {output_path}")

        # Step 2: Generate audio via Google TTS (Orus voice)
        await broadcast_log("🔊 Génération audio (Google TTS - Orus)...")
        audio_ok = await _generate_tts_audio(result, article_title, article_id)
        if audio_ok:
            await broadcast_log(f"✓ Audio généré")
        else:
            await broadcast_log("⚠ Audio échoué")

    except Exception as e:
        await broadcast_log(f"⚠ Erreur: {str(e)[:100]}")
    await broadcast_log("__done__")


async def _generate_tts_audio(text, title, article_id):
    """Generate MP3 from text using Google Gemini TTS API directly (Orus voice)."""
    try:
        import requests
        import base64
        from gnl_core.config import get_config
        from gnl_core.db import get_db

        config = get_config()
        audio_parts = config.get('AUDIO_PARTS_FOLDER', '')
        audio_dir = os.path.join(audio_parts, 'saved-articles', 'linkedin')
        os.makedirs(audio_dir, exist_ok=True)

        api_key = os.environ.get('GOOGLE_AI_API_KEY', '')
        if not api_key:
            # Try loading from .env
            from dotenv import load_dotenv
            load_dotenv()
            api_key = os.environ.get('GOOGLE_AI_API_KEY', '')
        if not api_key:
            return False

        tts_model = os.environ.get('TTS_MODEL', 'gemini-2.5-flash-preview-tts')
        url = f'https://generativelanguage.googleapis.com/v1beta/models/{tts_model}:generateContent?key={api_key}'

        payload = {
            'contents': [{'parts': [{'text': text}]}],
            'generationConfig': {
                'responseModalities': ['AUDIO'],
                'speechConfig': {
                    'voiceConfig': {
                        'prebuiltVoiceConfig': {'voiceName': os.environ.get('TTS_VOICE', 'Orus')}
                    }
                }
            }
        }

        loop = asyncio.get_event_loop()

        def _call_tts():
            resp = requests.post(url, json=payload, timeout=300)
            if resp.status_code != 200:
                return None
            data = resp.json()
            audio_b64 = data['candidates'][0]['content']['parts'][0]['inlineData']['data']
            return base64.b64decode(audio_b64)

        audio_bytes = await loop.run_in_executor(None, _call_tts)
        if not audio_bytes:
            return False

        safe_title = "".join(c if c.isalnum() or c in '-_ ' else '' for c in (title or f'article-{article_id}'))[:60]
        audio_path = os.path.join(audio_dir, f"{safe_title}.mp3")

        # Convert PCM (L16, 24kHz) to MP3 via ffmpeg
        import subprocess
        pcm_path = audio_path.replace('.mp3', '.pcm')
        with open(pcm_path, 'wb') as f:
            f.write(audio_bytes)
        subprocess.run([
            'ffmpeg', '-y', '-f', 's16le', '-ar', '24000', '-ac', '1',
            '-i', pcm_path, audio_path
        ], capture_output=True)
        os.unlink(pcm_path)

        with get_db() as conn:
            conn.execute("UPDATE saved_articles SET audio_path=? WHERE id=?", (audio_path, article_id))
            conn.commit()
        return True
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False


@app.post("/content/generate")
async def generate_content(request: Request):
    """Generate source PDF from AWS content."""
    form = await request.form()
    source = form.get("source")
    param = form.get("param")
    category = form.get("category", "")

    asyncio.create_task(_generate_content(source, param, category))
    return {"status": "started"}


async def _generate_content(source, param, category=""):
    loop = asyncio.get_event_loop()
    cat_label = f" [{category}]" if category else ""
    await broadcast_log(f"▶ Génération contenu: {source} ({param}){cat_label}")
    try:
        if source == "whats-new":
            from gnl_core.content import generate_whats_new

            def on_progress(msg):
                asyncio.run_coroutine_threadsafe(broadcast_log(msg), loop)

            # Support multiple months (comma-separated)
            months = [m.strip() for m in param.split(',') if m.strip()]
            result = await loop.run_in_executor(None, lambda: generate_whats_new(months, category_filter=category or None, on_progress=on_progress))
            if result:
                await broadcast_log(f"✓ PDF généré: {result}")
            else:
                await broadcast_log("⚠ Aucune annonce trouvée pour cette période")
    except Exception as e:
        await broadcast_log(f"⚠ Erreur: {str(e)[:100]}")
    await broadcast_log("__done__")


@app.post("/refresh")
async def refresh():
    """Force UI refresh."""
    await broadcast_status()
    return {"status": "ok"}


@app.post("/admin/save")
async def admin_save(request: Request):
    """Save configuration to gnl-config.json with validation."""
    from gnl_core.config import save_config
    form = await request.form()
    
    config_keys = ['AUDIO_PARTS_FOLDER', 'GNL_BACKLOG', 'PDF_PARTS_FOLDER', 
                   'NOTEBOOKLM_LANGUAGE', 'DEFAULT_SPEED', 'MCP_DOWNLOAD_TIMEOUT',
                   'MAX_GENERATION_RETRIES', 'TEST_MODE', 'TEST_GENERATION_DELAY',
                   'BEDROCK_MODEL_ID', 'AWS_REGION', 'AWS_PROFILE']
    
    data = {key: form.get(key, '') for key in config_keys}

    # Validation
    errors = []
    for path_key in ['AUDIO_PARTS_FOLDER', 'GNL_BACKLOG', 'PDF_PARTS_FOLDER']:
        if data[path_key] and not os.path.exists(data[path_key]):
            errors.append(f"{path_key}: chemin inexistant")
    try:
        speed = float(data['DEFAULT_SPEED'])
        if not (0.5 <= speed <= 3):
            errors.append("DEFAULT_SPEED: doit être entre 0.5 et 3")
    except ValueError:
        errors.append("DEFAULT_SPEED: nombre invalide")
    try:
        timeout = int(data['MCP_DOWNLOAD_TIMEOUT'])
        if timeout < 60:
            errors.append("MCP_DOWNLOAD_TIMEOUT: minimum 60 secondes")
    except ValueError:
        errors.append("MCP_DOWNLOAD_TIMEOUT: nombre invalide")
    try:
        retries = int(data['MAX_GENERATION_RETRIES'])
        if retries < 1:
            errors.append("MAX_GENERATION_RETRIES: minimum 1")
    except ValueError:
        errors.append("MAX_GENERATION_RETRIES: nombre invalide")

    if errors:
        await broadcast_log(f"⚠ Validation: {'; '.join(errors)}")
        return {"status": "error", "errors": errors}

    # Handle SCHEDULER fields from form before saving
    from gnl_core.config import get_config
    config = get_config()
    sched = config.get('SCHEDULER', {})
    for job_key in sched:
        time_field = f"SCHEDULER_{job_key}_time"
        enabled_field = f"SCHEDULER_{job_key}_enabled"
        if time_field in form.keys():
            sched[job_key]['time'] = form.get(time_field)
        sched[job_key]['enabled'] = enabled_field in form.keys()
    data['SCHEDULER'] = sched
    save_config(data)

    await broadcast_log("✓ Configuration sauvegardée")
    return {"status": "ok"}


@app.get("/admin/export")
async def admin_export():
    """Download config as JSON file."""
    from gnl_core.config import export_config
    from fastapi.responses import Response
    return Response(
        content=export_config(),
        media_type="application/octet-stream",
        headers={"Content-Disposition": "attachment; filename=gnl-config.json"}
    )


@app.post("/admin/import")
async def admin_import(request: Request):
    """Import config from uploaded JSON file."""
    from gnl_core.config import import_config
    form = await request.form()
    file = form.get("config_file")
    content = (await file.read()).decode()
    import_config(content)
    await broadcast_log("✓ Configuration importée")
    return {"status": "ok"}


@app.get("/api/quota-check/{parent_id}")
async def quota_check(parent_id: int):
    """Check if quota is sufficient for this edition."""
    from gnl_core.db import parent_status, get_db
    from datetime import datetime, timezone, timedelta

    s = parent_status(parent_id)
    to_generate = s['total'] - (s['generated'] or 0)

    if to_generate == 0:
        return {"sufficient": True, "to_generate": 0, "remaining": 0}

    remaining = _get_quota()

    sufficient = to_generate <= remaining
    pct = int(remaining / s['total'] * 100) if not sufficient else 100

    return {
        "sufficient": sufficient,
        "to_generate": to_generate,
        "remaining": remaining,
        "total": s['total'],
        "pct": pct
    }
@app.post("/action/{action}/{parent_id}")
async def run_action(action: str, parent_id: int, request: Request = None):
    asyncio.create_task(_run_action(action, parent_id))
    return {"status": "started", "action": action, "parent_id": parent_id}
async def _wait_quota_reset(deliver_start, deliver_timeout):
    """Wait until quota resets (midnight Pacific = 3:00 AM EST)."""
    import time as _time
    from datetime import datetime, timezone, timedelta
    pacific = timezone(timedelta(hours=-7))
    now = datetime.now(pacific)
    tomorrow = (now + timedelta(days=1)).replace(hour=0, minute=5, second=0, microsecond=0)
    wait_seconds = (tomorrow - now).total_seconds()
    # Cap wait to remaining timeout
    remaining = deliver_timeout - (_time.time() - deliver_start)
    wait_seconds = min(wait_seconds, remaining)
    if wait_seconds > 0:
        await asyncio.sleep(wait_seconds)


async def _run_action(action: str, parent_id: int):
    await broadcast_log(f"▶ {action} (parent_id={parent_id})")
    try:
        global _stop_signal
        _stop_signal = False
        loop = asyncio.get_event_loop()
        def _stopped():
            return _stop_signal
        if action == "generate":
            from gnl_core.generate import generate
            def on_progress(rec):
                asyncio.run_coroutine_threadsafe(broadcast_status(), loop)
            s, f = await loop.run_in_executor(None, lambda: generate(parent_id, on_progress=on_progress, should_stop=_stopped))
            await broadcast_log(f"✓ Generated {len(s)}, failed {len(f)}")
        elif action == "download":
            from gnl_core.download import download
            s, f = await loop.run_in_executor(None, lambda: download(parent_id))
            await broadcast_log(f"✓ Downloaded {len(s)}, failed {len(f)}")
        elif action == "convert":
            from gnl_core.convert import convert
            s, f = await loop.run_in_executor(None, lambda: convert(parent_id))
            await broadcast_log(f"✓ Converted {len(s)}, failed {len(f)}")
        elif action == "deliver":
            from gnl_core.generate import generate
            from gnl_core.download import download
            from gnl_core.convert import convert
            from gnl_core.db import parent_status, resolve_parent, get_db
            from gnl_core.combine import combine
            from datetime import datetime, timezone as tz, timedelta
            import time as _time

            deliver_timeout = int(os.environ.get('DELIVER_TIMEOUT', '48')) * 3600  # hours to seconds
            deliver_start = _time.time()

            while _time.time() - deliver_start < deliver_timeout:
                if _stop_signal:
                    await broadcast_log("⛔ Arrêté par l'utilisateur")
                    break

                s = parent_status(parent_id)

                # All done?
                if s['converted'] == s['total'] and s['combined'] > 0:
                    break

                # Generate phase (quota-aware)
                if s['generated'] < s['total']:
                    remaining_quota = _get_quota()
                    if remaining_quota > 0:
                        to_generate = s['total'] - s['generated']
                        max_count = min(to_generate, remaining_quota)
                        await broadcast_log(f"▶ GENERATE ({max_count} épisodes)")

                        def on_episode_generated(rec):
                            asyncio.run_coroutine_threadsafe(broadcast_status(), loop)

                        await loop.run_in_executor(None, lambda: generate(parent_id, max_count=max_count, on_progress=on_episode_generated, should_stop=_stopped))
                        await broadcast_status()
                    elif s['generated'] == 0:
                        await broadcast_log("⏳ Quota épuisé — attente du reset (minuit Pacific)...")
                        await _wait_quota_reset(deliver_start, deliver_timeout)
                        continue

                # Download phase
                s = parent_status(parent_id)
                if s['downloaded'] < s['generated']:
                    await broadcast_log(f"▶ DOWNLOAD ({s['generated'] - s['downloaded']} en attente)")

                    def on_dl_progress(rec):
                        asyncio.run_coroutine_threadsafe(broadcast_status(), loop)
                        asyncio.run_coroutine_threadsafe(broadcast_log(f"⬇ {rec['podcast_name']} téléchargé"), loop)

                    await loop.run_in_executor(None, lambda: download(parent_id, on_progress=on_dl_progress))
                    await broadcast_status()

                # Convert phase
                s = parent_status(parent_id)
                if s['downloaded'] > 0 and s['converted'] < s['downloaded']:
                    await broadcast_log("▶ CONVERT")
                    await loop.run_in_executor(None, lambda: convert(parent_id))
                    await broadcast_status()

                # Check if we need another loop iteration (quota was exhausted mid-generate)
                s = parent_status(parent_id)
                if s['generated'] < s['total']:
                    await broadcast_log("⏳ Quota épuisé — attente du reset (minuit Pacific)...")
                    await _wait_quota_reset(deliver_start, deliver_timeout)
                    continue

                # Not all downloaded yet — wait and retry
                if s['downloaded'] < s['generated']:
                    retry_delay = int(os.environ.get('DELIVER_RETRY_DELAY', '180'))
                    await broadcast_log(f"⏳ {s['downloaded']}/{s['generated']} téléchargés — nouvelle tentative dans {retry_delay//60} min...")
                    await asyncio.sleep(retry_delay)
                    continue

                # All generated and downloaded — combine
                if s['converted'] == s['total'] and s['combined'] == 0:
                    with get_db() as conn:
                        pf = conn.execute("SELECT parent_file FROM parent_configuration WHERE id=?", (parent_id,)).fetchone()['parent_file']
                    await broadcast_log("▶ COMBINE (complet)")
                    await broadcast_log("⏳ Combinaison en cours...")
                    await loop.run_in_executor(None, lambda: combine(parent_id, pf))
                    await broadcast_status()
                break

            s = parent_status(parent_id)
            await broadcast_log(f"✓ Deliver done: {s['converted']}/{s['total']} converted, combined={'✓' if s['combined'] else '✗'}")
        elif action == "clean":
            from gnl_core.clean import clean
            d, f = await loop.run_in_executor(None, lambda: clean(str(parent_id)))
            await broadcast_log(f"✓ Cleaned {d} notebooks")
            await broadcast_status()
        elif action == "reset":
            from gnl_core.clean import clean
            from gnl_core.db import get_db
            # Clean all notebooks
            d, f = clean("all")
            # Wipe DB
            with get_db() as conn:
                conn.execute("DELETE FROM podcast_download")
                conn.execute("DELETE FROM parent_configuration")
                conn.commit()
            await broadcast_log(f"🔄 Reset done")
            # Pre-fill the form will be handled client-side
    except Exception as e:
        await broadcast_log(f"⚠ Error: {str(e)[:100]}")
    finally:
        await broadcast_log("__done__")
        await broadcast_status()
@app.get("/content/preview/{theme}/{subtheme}/{filename}")
async def preview_prepare(theme: str, subtheme: str, filename: str):
    """Return PDF info for the prepare dialog."""
    from gnl_core.config import get_config
    import math
    from PyPDF2 import PdfReader
    config = get_config()
    inbox = config.get('INBOX_FOLDER', '')
    pdf_path = os.path.join(inbox, theme, subtheme, filename)
    if not os.path.isfile(pdf_path):
        return {"error": "File not found"}
    name = os.path.splitext(filename)[0]

    if theme == 'exams':
        # For exams: estimate questions from original file or formatted word
        from gnl_core.exams import get_exam_base
        base = get_exam_base(theme, subtheme)
        word_path = base / 'pdf-formatting' / 'word' / (name + '.docx')
        if word_path.exists():
            from docx import Document
            doc = Document(str(word_path))
            text = "\n".join([p.text for p in doc.paragraphs])
            import re
            total_questions = len(re.findall(r'Question\s+\d+:', text))
        else:
            # Estimate from original file
            if filename.lower().endswith('.docx'):
                from docx import Document
                doc = Document(pdf_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                import re
                total_questions = len(re.findall(r'Question\s+\d+:?', text))
                if total_questions == 0:
                    total_questions = max(1, len(doc.paragraphs) // 20)
            else:
                total_questions = 50  # rough estimate
        return {"name": name, "total_pages": total_questions, "suggested_pages": 5}

    if filename.lower().endswith('.docx'):
        from docx import Document
        doc = Document(pdf_path)
        total_pages = max(1, len(doc.paragraphs) // 30)
    else:
        total_pages = len(PdfReader(pdf_path).pages)
    suggested_pages = math.ceil(total_pages / 20)
    return {"name": name, "total_pages": total_pages, "suggested_pages": suggested_pages}


@app.post("/prepare-from-inbox")
async def prepare_from_inbox(request: Request):
    """Prepare a PDF already on disk (from content tab)."""
    form = await request.form()
    theme = form.get("theme", "")
    subtheme = form.get("subtheme", "")
    filename = form.get("filename", "")
    mode = form.get("mode", "pages")
    pages_per_episode = int(form.get("pages", 0))
    custom_name = form.get("name", "")
    pipeline = form.get("pipeline", "both")

    from gnl_core.config import get_config
    config = get_config()
    inbox = config.get('INBOX_FOLDER', '')
    pdf_path = os.path.join(inbox, theme, subtheme, filename)

    if not os.path.isfile(pdf_path):
        return {"status": "error", "error": "File not found"}

    name = custom_name or os.path.splitext(filename)[0]

    # Exam-specific: branched pipeline (trunk + anki/generate/both)
    if theme == 'exams':
        from gnl_core.exams import get_exam_base, step1_format, step2b_full_markdown, step3_highlight, step5_anki, split_exam_by_questions
        import shutil

        base = get_exam_base(theme, subtheme)
        loop = asyncio.get_event_loop()
        origin = 'dojo' if 'dojo' in filename.lower() else 'udemy'
        questions_per_chunk = pages_per_episode if pages_per_episode > 0 else 5
        do_anki = pipeline in ('anki', 'both', '')
        do_generate = pipeline in ('generate', 'both', '')

        def on_p(msg):
            asyncio.run_coroutine_threadsafe(broadcast_log(f"  {msg}"), loop)

        try:
            # === TRONC COMMUN (Format) ===
            # Step 1: Copy to origin + format
            origin_dir = base / 'pdf-formatting' / 'origin'
            origin_dir.mkdir(parents=True, exist_ok=True)
            origin_path = origin_dir / filename
            if not origin_path.exists():
                shutil.copy2(pdf_path, str(origin_path))

            await broadcast_log("▶ [FORMAT] Reformatage du document")
            word_path = await loop.run_in_executor(None, lambda: step1_format(str(origin_path), theme, subtheme, origin, on_progress=on_p))

            await broadcast_log("▶ [MARKDOWN] Word → Markdown")
            md_path = await loop.run_in_executor(None, lambda: step2b_full_markdown(word_path, theme, subtheme, on_progress=on_p))
            await broadcast_log(f"  ✓ {md_path}")

            # === BRANCHE ANKI ===
            if do_anki:
                await broadcast_log("▶ [HIGHLIGHT] Identification des réponses (Bedrock)")
                answers = await loop.run_in_executor(None, lambda: step3_highlight(md_path, on_progress=on_p))
                await broadcast_log(f"  ✓ {len(answers)} questions")

                await broadcast_log("▶ [ANKI] Génération .apkg")
                anki_path = await loop.run_in_executor(None, lambda: step5_anki(answers, md_path, theme, subtheme, on_progress=on_p))
                await broadcast_log(f"  ✓ {anki_path}")

                # Optional: generate draw.io diagrams
                from gnl_core.config import get_config as _gc
                if _gc().get('EXAM_GENERATE_DIAGRAMS', '0') == '1':
                    await broadcast_log("▶ [DIAGRAMS] Génération draw.io")
                    from gnl_core.diagrams import generate_exam_diagrams
                    diagrams_path = await loop.run_in_executor(None, lambda: generate_exam_diagrams(md_path, answers, theme, subtheme, on_progress=on_p))
                    await broadcast_log(f"  ✓ {diagrams_path}")

            # === BRANCHE GÉNÉRATION ===
            if do_generate:
                await broadcast_log(f"▶ [SPLIT] Découpage ({questions_per_chunk} questions/épisode)")
                result = await loop.run_in_executor(None, lambda: split_exam_by_questions(md_path, questions_per_chunk, name, theme, subtheme))
                await broadcast_log(f"  ✓ {len(result['files'])} morceaux")

                await broadcast_log("▶ [PRODUCTION] Insertion DB")
                from gnl_core.collect import collect
                from gnl_core.titles import generate_titles
                from gnl_core.db import get_db

                # Remove existing edition if re-processing
                with get_db() as conn:
                    existing = conn.execute(
                        "SELECT id FROM parent_configuration WHERE parent_file=? AND podcast_subtheme=?",
                        (name, subtheme)
                    ).fetchone()
                    if existing:
                        conn.execute("DELETE FROM podcast_download WHERE parent_configuration_id=?", (existing['id'],))
                        conn.execute("DELETE FROM parent_configuration WHERE id=?", (existing['id'],))
                        conn.commit()
                        await broadcast_log(f"  ♻ Ancienne édition supprimée (id={existing['id']})")

                parent_id = collect(result)
                generate_titles(parent_id)

            # Summary
            summary = "🎴 Anki" if do_anki and not do_generate else "📻 Production" if do_generate and not do_anki else "🎴 Anki + 📻 Production"
            await broadcast_log(f"✓ Pipeline terminé: {name} ({summary})")
            await broadcast_log("__done__")
            if do_generate:
                await broadcast_status()
                return {"status": "ok", "parent_id": parent_id}
            return {"status": "ok"}
        except Exception as e:
            await broadcast_log(f"⚠ Erreur: {str(e)[:100]}")
            return {"status": "error", "error": str(e)}

    # Standard flow for non-exam files
    # Calculate pages per episode if not provided
    if pages_per_episode <= 0:
        import math
        from PyPDF2 import PdfReader
        total_pages = len(PdfReader(pdf_path).pages)
        pages_per_episode = math.ceil(total_pages / 20)

    await broadcast_log(f"▶ Preparing {filename} (mode={mode}, {pages_per_episode} pages/épisode)")

    try:
        from gnl_core.split import split
        from gnl_core.collect import collect
        from gnl_core.titles import generate_titles

        result = split(pdf_path, pages_per_episode, name, podcast_theme=theme, podcast_subtheme=subtheme, mode=mode)
        parent_id = collect(result)
        count = generate_titles(parent_id)
        await broadcast_log(f"✓ Prepared: {len(result['files'])} chunks, parent_id={parent_id}, {count} titles")
        await broadcast_status()
        return {"status": "ok", "parent_id": parent_id}
    except Exception as e:
        await broadcast_log(f"⚠ Prepare failed: {str(e)[:100]}")
        return {"status": "error", "error": str(e)}


@app.post("/prepare")
async def prepare_pdf(request: Request):
    """Upload PDF, split, collect, generate titles."""
    from fastapi import UploadFile, Form
    import tempfile

    form = await request.form()
    pdf_file = form.get("pdf")
    pages = int(form.get("pages", 1))
    name = form.get("name", "")
    theme = form.get("theme", "")
    subtheme = form.get("subtheme", "")
    mode = form.get("mode", "pages")

    # Save uploaded file to temp
    tmp = os.path.join(tempfile.gettempdir(), pdf_file.filename)
    with open(tmp, "wb") as f:
        f.write(await pdf_file.read())

    await broadcast_log(f"▶ Preparing {pdf_file.filename} (mode={mode})")

    try:
        from gnl_core.split import split
        from gnl_core.collect import collect
        from gnl_core.titles import generate_titles

        result = split(tmp, pages, name, podcast_theme=theme, podcast_subtheme=subtheme, mode=mode)
        parent_id = collect(result)
        count = generate_titles(parent_id)
        os.unlink(tmp)
        await broadcast_log(f"✓ Prepared: {len(result['files'])} chunks, parent_id={parent_id}, {count} titles")
        await broadcast_status()
        return {"status": "ok", "parent_id": parent_id}
    except Exception as e:
        await broadcast_log(f"⚠ Prepare failed: {str(e)[:100]}")
        return {"status": "error", "error": str(e)}
@app.post("/toggle-test-mode")
async def toggle_test_mode():
    """Toggle TEST_MODE on/off."""
    current = os.environ.get('TEST_MODE', '0')
    new_val = '0' if current == '1' else '1'
    os.environ['TEST_MODE'] = new_val
    await broadcast_log(f"{'🧪 TEST MODE ON' if new_val == '1' else '🚀 TEST MODE OFF (real API)'}")
    return {"status": "ok", "test_mode": new_val == '1'}
@app.post("/schedule")
async def update_schedule(request: Request):
    """Update daily schedule time."""
    form = await request.form()
    new_time = form.get("time", "08:00")
    hour, minute = new_time.split(':')
    scheduler.reschedule_job('daily_deliver', trigger=CronTrigger(hour=int(hour), minute=int(minute)))
    os.environ['GNL_SCHEDULE_TIME'] = new_time
    return {"status": "updated", "time": new_time}
@app.websocket("/ws/logs")
async def websocket_logs(websocket: WebSocket):
    await websocket.accept()
    ws_clients.append(websocket)
    try:
        while True:
            await websocket.receive_text()
    except WebSocketDisconnect:
        ws_clients.remove(websocket)
