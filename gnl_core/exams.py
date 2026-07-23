"""Exam processing pipeline — from raw DOCX to Anki flashcards.

Structure:
    {EXAM_BASE}/pdf-formatting/origin/    ← DOCX brut (input)
    {EXAM_BASE}/pdf-formatting/word/      ← DOCX nettoyé (format)
    {EXAM_BASE}/pdf-formatting/pdf/       ← PDF converti (libreoffice)
    {EXAM_BASE}/pdf-formatting/compact-exam-versions/ ← PDF compact
    {EXAM_BASE}/Anki-generation/markdown/ ← Markdown compact
    {EXAM_BASE}/Anki-generation/anki/     ← Fichier Anki importable
"""

import os
import re
import shutil
import subprocess
from pathlib import Path


def get_exam_base(theme, subtheme):
    """Get base path for exam assets: INBOX_FOLDER/{theme}/{subtheme}/assets/"""
    inbox = os.environ.get('INBOX_FOLDER', '')
    return Path(inbox) / theme / subtheme / 'assets'


def step1_format(input_path, theme, subtheme, origin='udemy', on_progress=None):
    """Step 1: origin/ → word/ (clean DOCX)
    
    Args:
        input_path: Path to original DOCX in origin/
        origin: 'udemy' or 'dojo'
    Returns:
        Path to cleaned DOCX in word/
    """
    from docx import Document

    base = get_exam_base(theme, subtheme)
    word_dir = base / 'pdf-formatting' / 'word'
    word_dir.mkdir(parents=True, exist_ok=True)

    filename = Path(input_path).name
    output_path = word_dir / filename

    # Copy original to word folder
    shutil.copy2(input_path, output_path)

    # Process based on origin
    doc = Document(str(output_path))
    full_text = [para.text for para in doc.paragraphs]
    text = "\n".join(full_text)

    if origin == 'dojo':
        text = re.sub(r"References:.*?(?=Question|\Z)", "", text, flags=re.DOTALL | re.IGNORECASE)
    else:
        for p in [r"\[ \]", r"Ignoré.*?\n", r"Bonne réponse", r"Sélection correcte", r"Explication générale", r"via -.*?\n"]:
            text = re.sub(p, "", text)
        text = re.sub(r"\[Unofficial\].*?Tentative \d+\s*\n", "", text, flags=re.DOTALL)
        text = re.sub(r"Ressources\s*\nDomaine\s*\n.*?\n(?=Question)", "", text, flags=re.IGNORECASE)

    text = re.sub(r"\n\s*\n", "\n", text)
    text = re.sub(r"={50,}\n?", "", text)

    # Renumber questions and format options
    lines = text.split('\n')
    result_lines = []
    question_counter = 0
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if re.match(r'^\d+\.\s*Question$', stripped) or re.match(r'^Question$', stripped):
            question_counter += 1
            result_lines.append(f"Question {question_counter}:")
            i += 1
        elif re.match(r'^Question\s+\d+:?', stripped):
            question_counter += 1
            rest = re.sub(r'^Question\s+\d+:?\s*', '', stripped)
            result_lines.append(f"Question {question_counter}:")
            if rest:
                result_lines.append(rest)
            i += 1
        elif stripped == 'Incorrect' or re.match(r'^Correct\s+options?:', stripped, re.IGNORECASE):
            # Find the last question mark to identify where options start
            last_q_idx = -1
            for j in range(len(result_lines) - 1, -1, -1):
                if '?' in result_lines[j]:
                    last_q_idx = j
                    break

            if last_q_idx != -1:
                # Lines between question mark and here are options
                options = []
                for j in range(last_q_idx + 1, len(result_lines)):
                    if result_lines[j].strip():
                        options.append(result_lines[j].strip())

                # Remove option lines from result
                result_lines = result_lines[:last_q_idx + 1]

                # Add as bullet list
                for opt in options:
                    if not opt.startswith('- '):
                        result_lines.append(f"- {opt}")
                    else:
                        result_lines.append(opt)

            # Add Explanations marker
            result_lines.append("Explanations:")
            i += 1
        else:
            result_lines.append(line)
            i += 1

    # Save cleaned DOCX
    new_doc = Document()
    for line in result_lines:
        para = new_doc.add_paragraph()
        if re.match(r'^Question\s+\d+:', line.strip()):
            para.add_run(line).bold = True
        elif line.strip() == "Explanations:":
            para.add_run(line).bold = True
        else:
            para.add_run(line)
    new_doc.save(str(output_path))

    if on_progress:
        on_progress(f"origin/ → word/{filename}")

    return str(output_path)


def step2_convert_pdf(word_path, theme, subtheme, on_progress=None):
    """Step 2: word/ → pdf/ (LibreOffice conversion)
    
    Returns:
        Path to PDF
    """
    base = get_exam_base(theme, subtheme)
    pdf_dir = base / 'pdf-formatting' / 'pdf'
    pdf_dir.mkdir(parents=True, exist_ok=True)

    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(pdf_dir), word_path
    ], capture_output=True)

    pdf_path = pdf_dir / Path(word_path).with_suffix('.pdf').name

    if on_progress:
        on_progress(f"word/ → pdf/{pdf_path.name}")

    return str(pdf_path)


def step2b_full_markdown(word_path, theme, subtheme, on_progress=None):
    """Convert formatted DOCX to full Markdown (tronc commun step)."""
    from docx import Document

    base = get_exam_base(theme, subtheme)
    md_dir = base / 'pdf-formatting' / 'full-markdown'
    md_dir.mkdir(parents=True, exist_ok=True)

    name = Path(word_path).stem
    md_path = md_dir / f"{name}.md"

    doc = Document(word_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append('')
            continue

        # Detect question headers
        import re
        if re.match(r'^Question\s+\d+:', text):
            lines.append(f"## {text}")
        # Detect options (A) B) C) D) or A. B. C. D.)
        elif re.match(r'^[A-F][).]\s', text):
            lines.append(f"- **{text[0]}**) {text[2:].strip()}")
        else:
            lines.append(text)

    md_content = '\n'.join(lines)
    md_path.write_text(md_content, encoding='utf-8')

    if on_progress:
        on_progress(f"word/ → full-markdown/{md_path.name}")

    return str(md_path)


def step3_highlight(source_path, on_progress=None):
    """Step 3: Identify correct answers.
    
    Strategy (3-tier fallback):
      1. NotebookLM — create notebook, upload .md, query by batch of 10
      2. Bedrock (Claude) — if NLM fails or returns incomplete
      3. Regex — last resort for questions not resolved
    
    Args:
        source_path: Path to formatted DOCX or full Markdown file
    Returns:
        Dict {question_number: {"type": str, "options": list, "correct": list}}
    """
    import json

    # Read text
    if source_path.endswith('.md'):
        text = Path(source_path).read_text(encoding='utf-8')
        raw_text = text.replace('## ', '').replace('- **', '').replace('**)', ')')
    else:
        from docx import Document
        doc = Document(source_path)
        raw_text = "\n".join([para.text for para in doc.paragraphs])
        text = raw_text

    config_data = _get_config()

    # Split into question blocks
    questions = re.split(r'(Question\s+\d+:)', raw_text)
    question_blocks = []
    for i in range(1, len(questions), 2):
        if i + 1 < len(questions):
            q_num = re.search(r'\d+', questions[i])
            if q_num:
                question_blocks.append((q_num.group(), questions[i] + questions[i + 1]))

    # Load prompt template
    prompt_file = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))) / 'prompts' / 'exam-highlight.txt'
    if prompt_file.exists():
        prompt_template = prompt_file.read_text(encoding='utf-8')
    else:
        prompt_template = "Extract correct answers.\n\n{questions}\n\nReturn JSON."

    batch_size = int(config_data.get('EXAM_BATCH_SIZE', '10'))
    all_answers = {}

    # === TIER 1: NotebookLM ===
    nlm_success = False
    try:
        all_answers = _highlight_via_nlm(source_path, question_blocks, prompt_template, batch_size, config_data, on_progress)
        if len(all_answers) >= len(question_blocks) * 0.8:  # 80% success threshold
            nlm_success = True
            if on_progress:
                on_progress(f"✓ NotebookLM: {len(all_answers)}/{len(question_blocks)} questions")
    except Exception as e:
        if on_progress:
            on_progress(f"⚠ NotebookLM échoué: {str(e)[:60]} → fallback Bedrock")

    # === TIER 2: Bedrock (for missing questions) ===
    if not nlm_success or len(all_answers) < len(question_blocks):
        missing_blocks = [(num, content) for num, content in question_blocks if num not in all_answers]
        if missing_blocks:
            if on_progress:
                on_progress(f"Bedrock fallback: {len(missing_blocks)} questions manquantes")
            bedrock_answers = _highlight_via_bedrock(missing_blocks, prompt_template, batch_size, config_data, on_progress)
            all_answers.update(bedrock_answers)

    # === TIER 3: Regex fallback (last resort) ===
    still_missing = [(num, content) for num, content in question_blocks if num not in all_answers]
    if still_missing:
        if on_progress:
            on_progress(f"Regex fallback: {len(still_missing)} questions restantes")
        for num, content in still_missing:
            regex_result = _highlight_via_regex(num, content)
            if regex_result:
                all_answers[num] = regex_result

    return all_answers


def _highlight_via_nlm(source_path, question_blocks, prompt_template, batch_size, config_data, on_progress=None):
    """Tier 1: Use NotebookLM as knowledge base to identify answers.
    
    Creates a persistent notebook named {filename}-FULL with exams-default.txt
    as chat configuration (for future interactive audio generation).
    Queries use exam-highlight.txt prompt directly.
    """
    import json
    from notebooklm_tools.mcp.tools._utils import get_client
    from notebooklm_tools.services.notebooks import create_notebook, list_notebooks
    from notebooklm_tools.services.sources import add_source
    from notebooklm_tools.services.chat import query, configure_chat

    client = get_client()
    name = Path(source_path).stem
    notebook_title = f"{name}-FULL"

    # Check if notebook already exists
    notebook_id = None
    existing = list_notebooks(client)
    for nb in existing.get('notebooks', []):
        if nb.get('title') == notebook_title:
            notebook_id = nb['notebook_id']
            if on_progress:
                on_progress(f"Notebook existant: {notebook_title}")
            break

    # Create if not exists
    if not notebook_id:
        nb = create_notebook(client, title=notebook_title)
        notebook_id = nb['notebook_id']

        # Upload markdown as source
        add_source(client, notebook_id, source_type="file", file_path=source_path, wait=True)

        # Wait for indexation to complete (NLM needs time to process)
        import time
        time.sleep(10)

        # Configure chat with exams-default prompt (for future interactive audio)
        exams_prompt_file = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))) / 'prompts' / 'exams-default.txt'
        if exams_prompt_file.exists():
            exams_prompt = exams_prompt_file.read_text(encoding='utf-8')
            configure_chat(client, notebook_id, goal="custom", custom_prompt=exams_prompt)

        if on_progress:
            on_progress(f"Notebook créé: {notebook_title}")

    all_answers = {}

    # Query by batch using exam-highlight prompt (with retry)
    for batch_start in range(0, len(question_blocks), batch_size):
        batch = question_blocks[batch_start:batch_start + batch_size]
        batch_text = "\n\n".join([content for _, content in batch])
        query_text = prompt_template.replace('{questions}', batch_text)

        # Retry up to 3 times (indexation may still be in progress)
        answer_text = ''
        for attempt in range(3):
            try:
                result = query(client, notebook_id, query_text, timeout=120)
                answer_text = result.get('answer', '')
                break
            except Exception as e:
                if attempt < 2:
                    import time
                    time.sleep(5 * (attempt + 1))
                    if on_progress:
                        on_progress(f"Retry {attempt + 1}/3...")
                else:
                    raise

        # Parse JSON from response
        json_match = re.search(r'\{.*\}', answer_text, re.DOTALL)
        if json_match:
            try:
                parsed = json.loads(json_match.group())
                all_answers.update(parsed)
            except json.JSONDecodeError:
                # Try fixing common issues
                fixed = json_match.group().replace('\n', ' ')
                fixed = re.sub(r',\s*}', '}', fixed)
                fixed = re.sub(r',\s*]', ']', fixed)
                try:
                    parsed = json.loads(fixed)
                    all_answers.update(parsed)
                except json.JSONDecodeError:
                    pass

        if on_progress:
            on_progress(f"NLM batch {batch_start//batch_size + 1}: {len(batch)} questions")

    return all_answers


def _highlight_via_bedrock(question_blocks, prompt_template, batch_size, config_data, on_progress=None):
    """Tier 2: Use Bedrock/Claude for structured extraction."""
    import json
    import boto3
    from botocore.config import Config

    model_id = config_data.get('BEDROCK_MODEL_ID', 'us.anthropic.claude-sonnet-4-20250514-v1:0')
    region = config_data.get('AWS_REGION', 'ca-central-1')
    profile = config_data.get('AWS_PROFILE', '')
    max_tokens = int(config_data.get('BEDROCK_MAX_TOKENS', '4096'))

    session = boto3.Session(profile_name=profile, region_name=region)
    client = session.client('bedrock-runtime', config=Config(read_timeout=600))

    all_answers = {}

    def _process_batch(batch):
        batch_text = "\n\n".join([content for _, content in batch])
        prompt = prompt_template.replace('{questions}', batch_text)

        for attempt in range(3):
            try:
                response = client.converse(
                    modelId=model_id,
                    messages=[{"role": "user", "content": [{"text": prompt}]}],
                    inferenceConfig={"maxTokens": max_tokens}
                )
                result_text = response['output']['message']['content'][0]['text']
                json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
                if json_match:
                    raw_json = json_match.group()
                    try:
                        return json.loads(raw_json)
                    except json.JSONDecodeError:
                        fixed = raw_json.replace('\n', ' ')
                        fixed = re.sub(r',\s*}', '}', fixed)
                        fixed = re.sub(r',\s*]', ']', fixed)
                        try:
                            return json.loads(fixed)
                        except json.JSONDecodeError:
                            continue
            except Exception:
                if attempt == 2:
                    return {}
        return {}

    # Build and process batches in parallel
    from concurrent.futures import ThreadPoolExecutor, as_completed
    parallel = int(config_data.get('EXAM_PARALLEL_BATCHES', '3'))
    batches = []
    for i in range(0, len(question_blocks), batch_size):
        batches.append(question_blocks[i:i + batch_size])

    with ThreadPoolExecutor(max_workers=parallel) as executor:
        futures = {}
        for idx, batch in enumerate(batches):
            future = executor.submit(_process_batch, batch)
            futures[future] = idx

        for future in as_completed(futures):
            idx = futures[future]
            if on_progress:
                on_progress(f"Bedrock batch {idx + 1}/{len(batches)} ✓")
            result = future.result()
            if result:
                all_answers.update(result)

    return all_answers


def _highlight_via_regex(num, content):
    """Tier 3: Last resort regex-based extraction.
    Tries to find correct answer markers in the text (e.g., 'Correct', checkmarks, etc.)
    """
    options = re.findall(r'([A-F])[).]\s*(.+?)(?=\n[A-F][).]|\n\n|$)', content, re.DOTALL)
    if not options:
        return None

    option_texts = [f"{letter}) {text.strip().split(chr(10))[0]}" for letter, text in options]

    # Try to detect correct answers from common markers
    correct = []
    for letter, text in options:
        # Look for "Correct" or "✓" or "(correct)" near this option
        if re.search(r'(?i)\bcorrect\b|✓|✔', text):
            correct.append(f"{letter}) {text.strip().split(chr(10))[0]}")

    # If no correct found via markers, check explanation section
    if not correct:
        explanation_match = re.search(r'(?i)(?:explanation|answer|correct answer)[:\s]+([A-F](?:\s*,\s*[A-F])*)', content)
        if explanation_match:
            letters = re.findall(r'[A-F]', explanation_match.group(1))
            for letter in letters:
                for opt_letter, opt_text in options:
                    if opt_letter == letter:
                        correct.append(f"{opt_letter}) {opt_text.strip().split(chr(10))[0]}")

    if not correct:
        return None

    q_type = "multiple" if len(correct) > 1 else "single"
    if re.search(r'(?i)select\s+and\s+order|arrange|sequence', content):
        q_type = "order"

    return {"type": q_type, "options": option_texts, "correct": correct}


def step4_compact(source_path, answers, theme, subtheme, on_progress=None):
    """Step 4: Generate compact version (Markdown + PDF).
    Uses structured answers from step3 (type, options, correct).
    
    Args:
        source_path: Path to formatted DOCX or full Markdown
    Returns:
        Path to compact markdown
    """
    import markdown
    from weasyprint import HTML

    base = get_exam_base(theme, subtheme)
    name = Path(source_path).stem

    # Read text from DOCX or Markdown
    if source_path.endswith('.md'):
        text = Path(source_path).read_text(encoding='utf-8')
        text = text.replace('## ', '').replace('- **', '').replace('**)', ')')
    else:
        from docx import Document
        doc = Document(source_path)
        text = "\n".join([para.text for para in doc.paragraphs])

    # Extract question texts
    questions = re.split(r'(Question\s+\d+:)', text)
    question_texts = {}
    for i in range(1, len(questions), 2):
        if i + 1 < len(questions):
            q_num = re.search(r'\d+', questions[i])
            if q_num:
                num = q_num.group()
                lines = questions[i + 1].strip().split('\n')
                # Get question text (before options/explanation)
                q_text_lines = []
                for line in lines:
                    if line.strip().startswith('- ') or 'Explanation' in line or 'Hence,' in line or line.strip() == 'Incorrect':
                        break
                    if line.strip():
                        q_text_lines.append(line.strip())
                question_texts[num] = ' '.join(q_text_lines[:3])

    compact_lines = [f"# {name} — Compact Version\n"]

    for num in sorted(answers.keys(), key=lambda x: int(x)):
        entry = answers[num]
        
        # Handle both old format (list) and new format (dict)
        if isinstance(entry, list):
            # Old format: just a list of correct answers
            q_text = question_texts.get(num, '')
            compact_lines.append(f"\n**Question {num}:**")
            compact_lines.append("")
            compact_lines.append(q_text)
            compact_lines.append("")
            for ans in entry:
                compact_lines.append(f"- **{ans}**")
            continue

        # New structured format
        q_type = entry.get('type', 'single')
        options = entry.get('options', [])
        correct = entry.get('correct', [])
        q_text = question_texts.get(num, '')

        compact_lines.append(f"\n**Question {num}:** [{q_type}]")
        compact_lines.append("")
        compact_lines.append(q_text)
        compact_lines.append("")

        if q_type == 'order':
            # Show all options, then show correct order
            for opt in options:
                compact_lines.append(f"- {opt}")
            compact_lines.append("")
            compact_lines.append("**Correct order:**")
            for idx, item in enumerate(correct, 1):
                compact_lines.append(f"{idx}. **{item}**")
        else:
            # single, multiple, match — mark correct in bold
            correct_norm = [c.lower().replace('\u00a0', ' ')[:50] for c in correct]
            for opt in options:
                opt_norm = opt.lower().replace('\u00a0', ' ')[:50]
                is_correct = any(opt_norm in cn or cn in opt_norm for cn in correct_norm)
                if is_correct:
                    compact_lines.append(f"- **{opt}**")
                else:
                    compact_lines.append(f"- {opt}")

    # Save markdown
    md_dir = base / 'Anki-generation' / 'markdown'
    md_dir.mkdir(parents=True, exist_ok=True)
    md_path = md_dir / f"{name}.md"
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(compact_lines))

    # Generate compact PDF
    compact_pdf_dir = base / 'pdf-formatting' / 'compact-exam-versions'
    compact_pdf_dir.mkdir(parents=True, exist_ok=True)
    html_content = markdown.markdown('\n'.join(compact_lines))
    html_full = f"""<!DOCTYPE html><html><head><meta charset="utf-8">
    <style>body{{font-family:Arial;font-size:11px;margin:30px;}}
    h1{{color:#232f3e;border-bottom:2px solid #ff9900;}}
    strong{{color:#28a745;}}</style></head><body>{html_content}</body></html>"""
    HTML(string=html_full).write_pdf(str(compact_pdf_dir / f"{name}.pdf"))

    if on_progress:
        on_progress(f"compact → {md_path.name}")

    return str(md_path)


def step5_anki(compact_md_path, theme, subtheme, on_progress=None):
    """Step 5: Generate Anki .apkg package from compact markdown.
    Handles all question types: single, multiple, order, match.
    
    Returns:
        Path to .apkg file
    """
    import genanki
    import random

    base = get_exam_base(theme, subtheme)
    name = Path(compact_md_path).stem

    with open(compact_md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Model for exam cards
    font_size = os.environ.get('ANKI_FONT_SIZE', '16')
    model_id = random.randrange(1 << 30, 1 << 31)
    model = genanki.Model(
        model_id,
        f'{name} Model',
        fields=[
            {'name': 'Front'},
            {'name': 'Back'},
        ],
        templates=[{
            'name': 'Card 1',
            'qfmt': '{{Front}}',
            'afmt': '{{Back}}',
        }],
        css=f"""
            .card {{ text-align: left; font-family: Arial; font-size: {font_size}px; padding: 10px; }}
            ul {{ padding-left: 20px; }}
            ol {{ padding-left: 20px; }}
            li {{ margin-bottom: 5px; }}
            .correct {{ color: #28a745; font-weight: bold; }}
            b {{ color: #0073bb; }}
        """
    )

    # Deck with exam name
    deck_id = random.randrange(1 << 30, 1 << 31)
    deck = genanki.Deck(deck_id, name)

    # Parse questions from markdown
    questions = re.split(r'(\*\*Question\s+\d+:\*\*(?:\s*\[[a-z]+\])?)', content)

    cards_count = 0
    for i in range(1, len(questions), 2):
        q_header_raw = questions[i]
        q_header = re.sub(r'\*\*', '', q_header_raw).strip()
        q_type_match = re.search(r'\[(\w+)\]', q_header)
        q_type = q_type_match.group(1) if q_type_match else 'single'
        q_header_clean = re.sub(r'\s*\[\w+\]', '', q_header)

        q_body = questions[i + 1].strip() if i + 1 < len(questions) else ''
        lines = q_body.split('\n')

        # Get question text (first non-empty line)
        question_text = ''
        for l in lines:
            if l.strip() and not l.strip().startswith('- ') and not l.strip().startswith('**Correct') and 'correct order' not in l.lower():
                question_text = l.strip()
                break

        if q_type == 'order':
            # Options = regular list items, Correct order = numbered bold items
            options = []
            correct_order = []
            in_correct = False
            for line in lines:
                line = line.strip()
                if 'correct order' in line.lower():
                    in_correct = True
                elif in_correct and re.match(r'^\d+\.', line):
                    item = re.sub(r'^\d+\.\s*\*\*(.+)\*\*$', r'\1', line)
                    correct_order.append(item)
                elif line.startswith('- '):
                    options.append(line[2:])

            if question_text and options:
                options_html = "".join(f"<li>{o}</li>" for o in options)
                front = f"<b>{q_header_clean}</b><br><br>{question_text}<br><br><ul>{options_html}</ul>"

                order_html = "".join(f"<li><span class='correct'>{o}</span></li>" for o in correct_order)
                back = f"<b>{q_header_clean}</b><br><br>{question_text}<br><br><b>Correct order:</b><ol>{order_html}</ol>"

                note = genanki.Note(model=model, fields=[front, back])
                deck.add_note(note)
                cards_count += 1
        else:
            # single, multiple, match
            options = []
            for line in lines:
                line = line.strip()
                if line.startswith('- **') and line.endswith('**'):
                    opt_text = line[4:-2]
                    options.append(('correct', opt_text))
                elif line.startswith('- '):
                    opt_text = line[2:]
                    options.append(('wrong', opt_text))

            if question_text and options:
                options_html = "".join(f"<li>{opt_text}</li>" for _, opt_text in options)
                front = f"<b>{q_header_clean}</b><br><br>{question_text}<br><br><ul>{options_html}</ul>"

                back_items = []
                for status, opt_text in options:
                    if status == 'correct':
                        back_items.append(f"<li><span class='correct'>{opt_text}</span></li>")
                    else:
                        back_items.append(f"<li>{opt_text}</li>")
                back = f"<b>{q_header_clean}</b><br><br>{question_text}<br><br><ul>{''.join(back_items)}</ul>"

                note = genanki.Note(model=model, fields=[front, back])
                deck.add_note(note)
                cards_count += 1

    # Save .apkg
    anki_dir = base / 'Anki-generation' / 'anki'
    anki_dir.mkdir(parents=True, exist_ok=True)
    apkg_path = anki_dir / f"{name}.apkg"
    genanki.Package(deck).write_to_file(str(apkg_path))

    if on_progress:
        on_progress(f"anki → {apkg_path.name} ({cards_count} cards)")

    return str(apkg_path)


def _get_config():
    from gnl_core.config import get_config
    return get_config()


def split_exam_by_questions(md_path, questions_per_chunk=5, name=None, theme='exams', subtheme='sap-c02'):
    """Split a full Markdown exam by questions into .md chunks.
    
    Args:
        md_path: Path to the full Markdown (from assets/pdf-formatting/full-markdown/)
        questions_per_chunk: Number of questions per split
        name: Edition name
        theme, subtheme: For path resolution
    Returns:
        Split result dict compatible with collect()
    """

    md_content = Path(md_path).read_text(encoding='utf-8')
    
    if not name:
        name = Path(md_path).stem

    # Split by "## Question N:" headers
    parts = re.split(r'(## Question\s+\d+:)', md_content)
    
    # Rebuild question blocks
    question_blocks = []
    for i in range(1, len(parts), 2):
        if i + 1 < len(parts):
            question_blocks.append(parts[i] + parts[i + 1])

    # Group into chunks
    chunks = []
    for i in range(0, len(question_blocks), questions_per_chunk):
        chunk = question_blocks[i:i + questions_per_chunk]
        chunks.append("\n\n".join(chunk))

    # Output directory
    pdf_parts = os.environ.get('PDF_PARTS_FOLDER', '')
    output_dir = Path(pdf_parts) / theme / subtheme / name
    output_dir.mkdir(parents=True, exist_ok=True)

    # Write each chunk as .md
    files_list = []
    for idx, chunk_text in enumerate(chunks, 1):
        chunk_path = output_dir / f"p{idx}.md"
        chunk_path.write_text(chunk_text.strip(), encoding='utf-8')

        files_list.append({
            'fullPath': str(chunk_path),
            'parentDir': name,
            'fileName': f'p{idx}.md',
            'sourceType': 'LocalStorage',
            'podcastTheme': theme,
            'podcastSubfolder': subtheme,
        })

    return {
        'splitConfiguration': f'{len(chunks)}ck-{questions_per_chunk}q',
        'files': files_list
    }
