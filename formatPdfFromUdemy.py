from docx import Document
import re
import sys
import subprocess
import os
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()


def clean_dojo_document(input_path, output_path):
    """Remove References sections from Dojo documents."""
    doc = Document(input_path)
    
    # Extract all text
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    text = "\n".join(full_text)
    
    # Remove References sections (from "References:" to next question pattern)
    text = re.sub(r"References:.*?(?=Question|\Z)", "", text, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove multiple blank lines
    text = re.sub(r"\n\s*\n", "\n", text)
    
    # Renumber all questions sequentially (handles both patterns and removes duplicates)
    lines = text.split('\n')
    result_lines = []
    question_counter = 0
    seen_questions = set()
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Check if line matches "N. Question" pattern
        if re.match(r'^\d+\.\s*Question$', stripped):
            question_counter += 1
            result_lines.append(f"Question {question_counter}:")
            i += 1
        # Check if line matches standalone "Question" pattern
        elif re.match(r'^Question$', stripped):
            question_counter += 1
            result_lines.append(f"Question {question_counter}:")
            i += 1
        # Check if line is already formatted "Question N"
        elif re.match(r'^Question\s+\d+:?$', stripped):
            # Skip duplicates
            if stripped not in seen_questions:
                question_counter += 1
                result_lines.append(f"Question {question_counter}:")
                seen_questions.add(stripped)
            i += 1
        # Check if line contains "Select and order" pattern
        elif re.search(r'Select and order', stripped, re.IGNORECASE):
            result_lines.append(line)
            i += 1
            
            # Collect options until "Incorrect"
            options = []
            seen_options = set()
            while i < len(lines) and not lines[i].strip().startswith("Incorrect"):
                if lines[i].strip():
                    option_text = lines[i].strip()
                    # Only add if not a duplicate
                    if option_text not in seen_options:
                        options.append(option_text)
                        seen_options.add(option_text)
                i += 1
            
            # Add options as list
            for option in options:
                result_lines.append(f"- {option}")
            
            # Add "Explanations:" title
            result_lines.append("Explanations:")
            
            # Skip the "Incorrect" line
            if i < len(lines) and lines[i].strip().startswith("Incorrect"):
                i += 1
        # Check if line is "Incorrect" - format options before it
        elif stripped.startswith("Incorrect"):
            # Find the last question mark before this line
            last_question_idx = -1
            for j in range(len(result_lines) - 1, -1, -1):
                if '?' in result_lines[j]:
                    last_question_idx = j
                    break
            
            if last_question_idx != -1:
                # Extract options (lines between question and "Incorrect")
                options = []
                for j in range(last_question_idx + 1, len(result_lines)):
                    if result_lines[j].strip():
                        options.append(result_lines[j].strip())
                
                # Remove the option lines from result
                result_lines = result_lines[:last_question_idx + 1]
                
                # Add options as list items
                for option in options:
                    result_lines.append(f"- {option}")
                
                # Add "Explanations:" title
                result_lines.append("Explanations:")
            
            # Skip the "Incorrect" line
            i += 1
        else:
            result_lines.append(line)
            i += 1
    
    text = '\n'.join(result_lines)
    
    # Use Bedrock to extract correct answers for each question
    print("Extracting correct answers using Bedrock...")
    lines_list = text.split('\n')
    correct_answers_map = {}  # Use highlight_correct_options.py instead
    print(f"Bedrock extracted answers for {len(correct_answers_map)} questions")
    
    # Mark options that match correct answers as bold
    marked_lines = []
    current_question = None
    in_select_and_order = False
    
    for line in lines_list:
        # Check if we're in a "Select and order" section
        if re.search(r'Select and order', line, re.IGNORECASE):
            in_select_and_order = True
            marked_lines.append(("normal", line))
        # Track current question
        elif re.match(r'^Question\s+\d+:', line.strip()):
            q_match = re.search(r'Question\s+(\d+):', line)
            if q_match:
                current_question = q_match.group(1)
            in_select_and_order = False
            marked_lines.append(("normal", line))
        # Check if this is an option line (starts with "- ")
        elif line.strip().startswith("- "):
            option_text = line.strip()[2:]  # Remove "- " prefix
            
            # Check if this option is in the correct answers for current question
            matched = False
            if current_question and current_question in correct_answers_map:
                correct_answers = correct_answers_map[current_question]
                
                for correct_ans in correct_answers:
                    # Normalize for comparison
                    opt_normalized = ' '.join(option_text.split()).lower()
                    ans_normalized = ' '.join(correct_ans.split()).lower()
                    
                    # Check for match
                    if opt_normalized == ans_normalized or \
                       (len(opt_normalized) > 50 and len(ans_normalized) > 50 and \
                        (opt_normalized in ans_normalized or ans_normalized in opt_normalized) and \
                        min(len(opt_normalized), len(ans_normalized)) / max(len(opt_normalized), len(ans_normalized)) > 0.8):
                        matched = True
                        break
            
            if matched:
                marked_lines.append(("option_bold", line))
            else:
                marked_lines.append(("option", line))
        else:
            # Reset flag when we leave the select and order section
            if line.strip() and not line.strip().startswith("- "):
                in_select_and_order = False
            marked_lines.append(("normal", line))
    
    # Save to new document
    new_doc = Document()
    for line_type, line in marked_lines:
        para = new_doc.add_paragraph()
        # Check if line is a question header
        if re.match(r'^Question\s+\d+:$', line.strip()):
            run = para.add_run(line)
            run.bold = True
        # Check if line is "Explanations:"
        elif line.strip() == "Explanations:":
            run = para.add_run(line)
            run.bold = True
        # Check if line contains "Hence, the correct answer" pattern
        elif re.search(r'Hence, the correct answers? (?:are|is):', line, re.IGNORECASE):
            # Split and format only the pattern
            match = re.search(r'(Hence, the correct answers? (?:are|is):)', line, re.IGNORECASE)
            if match:
                before = line[:match.start()]
                pattern = match.group(1)
                after = line[match.end():]
                
                if before:
                    para.add_run(before)
                run = para.add_run(pattern)
                run.bold = True
                run.underline = True
                if after:
                    para.add_run(after)
            else:
                para.add_run(line)
        # Check if this is a bold option
        elif line_type == "option_bold":
            run = para.add_run(line)
            run.bold = True
        else:
            para.add_run(line)
    
    new_doc.save(output_path)


def clean_word_for_anki(input_path, output_path):
    doc = Document(input_path)
    filename = Path(input_path).stem
    
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    text = "\n".join(full_text)

    # Remove Udemy-specific patterns
    patterns_to_remove = [
        r"\[ \]", 
        r"Ignoré.*?\n",
        r"Bonne réponse", 
        r"Sélection correcte", 
        r"Explication générale",
        r"via -.*?\n"
    ]
    
    for p in patterns_to_remove:
        text = re.sub(p, "", text)

    # Replace title section with filename
    text = re.sub(r"\[Unofficial\].*?Tentative \d+\s*\n", filename + "\n", text, flags=re.DOTALL)
    
    # Remove References sections (from Dojo logic)
    text = re.sub(r"References?:.*?(?=Question|\Z)", "", text, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove Ressources/Domaine sections
    text = re.sub(r"Ressources\s*\nDomaine\s*\n.*?\n(?=Question)", "", text, flags=re.IGNORECASE)

    # Remove multiple blank lines
    text = re.sub(r"\n\s*\n", "\n", text)
    text = re.sub(r"={50,}\n?", "", text)

    # Renumber questions sequentially (from Dojo logic)
    lines = text.split('\n')
    result_lines = []
    question_counter = 0
    seen_questions = set()
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        if re.match(r'^\d+\.\s*Question$', stripped):
            question_counter += 1
            result_lines.append(f"Question {question_counter}:")
            i += 1
        elif re.match(r'^Question$', stripped):
            question_counter += 1
            result_lines.append(f"Question {question_counter}:")
            i += 1
        elif re.match(r'^Question\s+\d+', stripped):
            # Matches "Question N" with or without colon, with or without text after
            question_counter += 1
            result_lines.append(f"Question {question_counter}:")
            # Extract text after "Question N:" or "Question N"
            rest = re.sub(r'^Question\s+\d+:?\s*', '', stripped, flags=re.IGNORECASE)
            if rest:
                result_lines.append(rest)
            i += 1
        elif re.search(r'Select and order', stripped, re.IGNORECASE):
            result_lines.append(line)
            i += 1
            
            options = []
            seen_options = set()
            while i < len(lines) and not lines[i].strip().startswith("Incorrect"):
                if lines[i].strip():
                    option_text = lines[i].strip()
                    if option_text not in seen_options:
                        options.append(option_text)
                        seen_options.add(option_text)
                i += 1
            
            for option in options:
                result_lines.append(f"- {option}")
            
            result_lines.append("Explanations:")
            
            if i < len(lines) and lines[i].strip().startswith("Incorrect"):
                i += 1
        elif re.match(r'^Correct\s+options?:', stripped, re.IGNORECASE):
            # Find last question mark before this line
            last_question_idx = -1
            for j in range(len(result_lines) - 1, -1, -1):
                if '?' in result_lines[j]:
                    last_question_idx = j
                    break
            
            if last_question_idx != -1:
                # Extract options between question mark and this line
                options = []
                for j in range(last_question_idx + 1, len(result_lines)):
                    if result_lines[j].strip():
                        options.append(result_lines[j].strip())
                
                # Remove option lines
                result_lines = result_lines[:last_question_idx + 1]
                
                # Add as bullet list
                for option in options:
                    if not option.startswith('- '):
                        result_lines.append(f"- {option}")
                    else:
                        result_lines.append(option)
            
            # Keep the Correct option line
            result_lines.append(line)
            i += 1
        else:
            result_lines.append(line)
            i += 1
    
    text = '\n'.join(result_lines)

    # Save to new document
    new_doc = Document()
    for i, line in enumerate(text.split('\n')):
        para = new_doc.add_paragraph()
        if i == 0:
            run = para.add_run(line)
            run.bold = True
            para.alignment = 1
        elif re.match(r'^Question\s+\d+:', line.strip()):
            para.add_run(line).bold = True
        elif line.strip() == "Explanations:":
            para.add_run(line).bold = True
        elif re.match(r'^(Correct|Incorrect)\s+options?:', line.strip(), re.IGNORECASE):
            run = para.add_run(line)
            run.bold = True
        elif re.search(r'Hence, the correct answers? (?:are|is):', line, re.IGNORECASE):
            match = re.search(r'(Hence, the correct answers? (?:are|is):)', line, re.IGNORECASE)
            if match:
                before = line[:match.start()]
                pattern = match.group(1)
                after = line[match.end():]
                
                if before:
                    para.add_run(before)
                run = para.add_run(pattern)
                run.bold = True
                run.underline = True
                if after:
                    para.add_run(after)
            else:
                para.add_run(line)
        else:
            para.add_run(line)
    
    new_doc.save(output_path)
    print(f"Traitement terminé ! Fichier enregistré sous : {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python formatPdfFromUdemy.py <filename> <origin>")
        print("  origin: 'udemy' or 'dojo'")
        sys.exit(1)
    
    filename = sys.argv[1]
    origin = sys.argv[2].lower()
    
    if origin not in ['udemy', 'dojo']:
        print("Error: origin must be 'udemy' or 'dojo'")
        sys.exit(1)
    
    # Get GNL_PROCESSING_PATH from environment
    gnl_processing_path = os.getenv('GNL_PROCESSING_PATH')
    if not gnl_processing_path:
        raise ValueError("GNL_PROCESSING_PATH not found in .env file")
    
    base_path = Path(gnl_processing_path).parent
    
    # Build input path from origin folder
    origin_file = base_path / "pdf-formatting" / "origin" / f"{filename}.docx"
    
    if not origin_file.exists():
        print(f"Error: File not found: {origin_file}")
        sys.exit(1)
    
    # Copy to word folder for processing
    word_folder = base_path / "pdf-formatting" / "word"
    word_folder.mkdir(parents=True, exist_ok=True)
    input_file = word_folder / f"{filename}.docx"
    
    import shutil
    shutil.copy2(origin_file, input_file)
    print(f"Copied from origin folder: {origin_file}")
    print(f"Processing copy: {input_file}")
    
    # Process document based on origin
    if origin == 'udemy':
        # Create temporary output file
        temp_file = input_file.with_suffix('.tmp.docx')
        clean_word_for_anki(str(input_file), str(temp_file))
        temp_file.replace(input_file)
        print("Document cleaned for Udemy format")
    else:
        # Clean Dojo document
        temp_file = input_file.with_suffix('.tmp.docx')
        clean_dojo_document(str(input_file), str(temp_file))
        temp_file.replace(input_file)
        print("Document cleaned for Dojo format (References removed)")
    
    # Convert to PDF
    exam_folder = Path(gnl_processing_path) / ".." / ".." / "exam"
    exam_folder.mkdir(parents=True, exist_ok=True)
    
    # Use LibreOffice to convert to PDF directly to exam folder
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(exam_folder), str(input_file)
    ], check=True)
    
    pdf_file = exam_folder / input_file.with_suffix('.pdf').name
    print(f"PDF saved to: {pdf_file}")