#!/usr/bin/env python3
"""Highlight correct options in exam documents using Bedrock AI."""

import fire
import re
import os
import json
import requests
import subprocess
import shutil
from pathlib import Path
from docx import Document
from dotenv import load_dotenv

load_dotenv()


def extract_correct_answers_with_bedrock(text):
    """Use Bedrock to extract correct answers for each question."""
    model_id = os.getenv('MOEDL_INFERENCE_ID', 'global.anthropic.claude-opus-4-5-20251101-v1:0')
    api_key = os.getenv('AWS_BEARER_TOKEN_BEDROCK', '')
    aws_region = os.getenv('AWS_REGION', 'us-east-1')
    
    if not api_key:
        print("⚠ Warning: AWS_BEARER_TOKEN_BEDROCK not found in .env file")
        return {}
    
    print(f"Using Bedrock model: {model_id}")
    
    # Split text into questions
    questions = re.split(r'(Question\s+\d+:)', text)
    question_blocks = []
    
    for i in range(1, len(questions), 2):
        if i + 1 < len(questions):
            q_num = re.search(r'\d+', questions[i])
            if q_num:
                content = questions[i] + questions[i + 1]
                question_blocks.append((q_num.group(), content))
    
    print(f"Found {len(question_blocks)} questions to process")
    
    all_answers = {}
    
    # Process in batches of 10
    for batch_start in range(0, len(question_blocks), 10):
        batch_end = min(batch_start + 10, len(question_blocks))
        batch = question_blocks[batch_start:batch_end]
        
        print(f"Processing questions {batch_start + 1}-{batch_end}...")
        
        batch_text = "\n\n".join([content for _, content in batch])
        
        prompt = f"""Extract the correct answer(s) for each question below.

{batch_text}

For each question, find the correct answer by looking for:
- "Hence, the correct answer is: ..." 
- "Hence, the correct answers are:" followed by "–" lines

Return ONLY a JSON object mapping question numbers to arrays of exact option texts:
{{"1": ["exact option text"], "2": ["option 1", "option 2"]}}

Return the exact text as it appears after "- " in the options list.
"""
        
        try:
            payload = {
                "messages": [{"role": "user", "content": [{"text": prompt}]}],
                "inferenceConfig": {"maxTokens": 4000, "temperature": 0.1}
            }
            
            url = f"https://bedrock-runtime.{aws_region}.amazonaws.com/model/{model_id}/converse"
            headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
            
            response = requests.post(url, json=payload, headers=headers)
            
            if response.status_code == 200:
                result = response.json()
                extracted = result['output']['message']['content'][0]['text']
                
                json_match = re.search(r'\{.*\}', extracted, re.DOTALL)
                if json_match:
                    batch_answers = json.loads(json_match.group())
                    all_answers.update(batch_answers)
                    print(f"  ✓ Extracted {len(batch_answers)} answers")
                else:
                    print(f"  ⚠ No JSON in response")
            else:
                print(f"  ✗ API error: {response.status_code}")
        except Exception as e:
            print(f"  ✗ Error: {str(e)}")
    
    print(f"✓ Total: {len(all_answers)} questions processed")
    return all_answers


def review_and_fix_document(word_file):
    """Review document and fix questions with no bold options."""
    doc = Document(word_file)
    
    # Track questions and their options
    questions_to_fix = []
    current_question = None
    current_question_text = []
    current_options = []
    has_bold_option = False
    in_question = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check if this is a question header
        q_match = re.match(r'^Question\s+(\d+):', text)
        if q_match:
            # Save previous question if it had no bold options
            if current_question and not has_bold_option and current_options:
                questions_to_fix.append({
                    'number': current_question,
                    'text': '\n'.join(current_question_text),
                    'options': current_options
                })
            
            # Start new question
            current_question = q_match.group(1)
            current_question_text = [text]
            current_options = []
            has_bold_option = False
            in_question = True
            continue
        
        if in_question:
            # Always collect text until next question
            current_question_text.append(text)
            
            # Check if this is an option
            if text.startswith('- '):
                is_bold = any(run.bold and len(run.text.strip()) > 0 for run in para.runs)
                current_options.append((para, text[2:].strip(), is_bold))
                if is_bold:
                    has_bold_option = True
    
    # Check last question
    if current_question and not has_bold_option and current_options:
        questions_to_fix.append({
            'number': current_question,
            'text': '\n'.join(current_question_text),
            'options': current_options
        })
    
    if not questions_to_fix:
        print("✓ All questions have at least one bold option")
        return
    
    print(f"⚠ Found {len(questions_to_fix)} questions with no bold options")
    print(f"  Questions: {', '.join([q['number'] for q in questions_to_fix])}")
    
    # Re-process these questions with Bedrock
    for question in questions_to_fix:
        print(f"  Re-processing Question {question['number']}...")
        
        question_text = question['text']
        correct_answers = extract_correct_answers_with_bedrock(question_text)
        
        if question['number'] in correct_answers:
            correct_options = correct_answers[question['number']]
            print(f"    Bedrock returned {len(correct_options)} correct options")
            
            # Apply bold to matching options
            fixed = False
            for para, option_text, _ in question['options']:
                for correct_ans in correct_options:
                    opt_normalized = ' '.join(option_text.split()).lower()
                    ans_normalized = ' '.join(correct_ans.split()).lower()
                    
                    if opt_normalized == ans_normalized or \
                       (len(opt_normalized) > 50 and len(ans_normalized) > 50 and \
                        (opt_normalized in ans_normalized or ans_normalized in opt_normalized) and \
                        min(len(opt_normalized), len(ans_normalized)) / max(len(opt_normalized), len(ans_normalized)) > 0.8):
                        for run in para.runs:
                            run.bold = True
                        fixed = True
                        print(f"    ✓ Fixed Question {question['number']}")
                        break
                if fixed:
                    break
            
            if not fixed:
                print(f"    ⚠ Could not match options for Question {question['number']}")
        else:
            print(f"    ⚠ No answer found by Bedrock for Question {question['number']}")
    
    # Save fixed document
    doc.save(word_file)
    print(f"✓ Fixed {len(questions_to_fix)} questions")


def main(filename: str, origin: str):
    """Highlight correct options in exam document.
    
    Args:
        filename: Name of the file without extension
        origin: 'udemy' or 'dojo'
    """
    if origin.lower() not in ['udemy', 'dojo']:
        raise ValueError("origin must be 'udemy' or 'dojo'")
    
    origin = origin.lower()
    
    # Get paths
    gnl_processing_path = os.getenv('GNL_PROCESSING_PATH')
    if not gnl_processing_path:
        raise ValueError("GNL_PROCESSING_PATH not found in .env file")
    
    base_path = Path(gnl_processing_path).parent
    word_file = base_path / "pdf-formatting" / "word" / f"{filename}.docx"
    
    if not word_file.exists():
        raise FileNotFoundError(f"Word file not found: {word_file}")
    
    print(f"Processing: {word_file}")
    
    # Read document
    doc = Document(str(word_file))
    text = '\n'.join([p.text for p in doc.paragraphs])
    
    # Extract correct answers using Bedrock
    correct_answers_map = extract_correct_answers_with_bedrock(text)
    
    if not correct_answers_map:
        print("⚠ No correct answers extracted, skipping highlighting")
        return
    
    # Apply bold formatting to correct options
    current_question = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Track current question
        q_match = re.match(r'^Question\s+(\d+):', text)
        if q_match:
            current_question = q_match.group(1)
            continue
        
        # Check if this is an option
        if text.startswith('- ') and current_question and current_question in correct_answers_map:
            option_text = text[2:].strip()
            correct_answers = correct_answers_map[current_question]
            
            # Check if this option matches any correct answer
            for correct_ans in correct_answers:
                opt_normalized = ' '.join(option_text.split()).lower()
                ans_normalized = ' '.join(correct_ans.split()).lower()
                
                if opt_normalized == ans_normalized or \
                   (len(opt_normalized) > 50 and len(ans_normalized) > 50 and \
                    (opt_normalized in ans_normalized or ans_normalized in opt_normalized) and \
                    min(len(opt_normalized), len(ans_normalized)) / max(len(opt_normalized), len(ans_normalized)) > 0.8):
                    # Make this paragraph bold
                    for run in para.runs:
                        run.bold = True
                    print(f"  ✓ Bolded option in Question {current_question}")
                    break
    
    # Save document
    doc.save(str(word_file))
    print(f"\n✓ Document updated with highlighted correct options")
    
    # Review and fix questions with no bold options
    print(f"\nReviewing document for missed questions...")
    review_and_fix_document(str(word_file))
    
    # Convert to PDF
    print(f"\nConverting to PDF...")
    pdf_folder = base_path / "pdf-formatting" / "pdf"
    pdf_folder.mkdir(parents=True, exist_ok=True)
    
    subprocess.run([
        "libreoffice", "--headless", "--convert-to", "pdf",
        "--outdir", str(pdf_folder), str(word_file)
    ], check=True)
    
    pdf_file = pdf_folder / f"{filename}.pdf"
    print(f"✓ PDF saved to: {pdf_file}")


if __name__ == "__main__":
    fire.Fire(main)
